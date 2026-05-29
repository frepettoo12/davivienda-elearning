#!/usr/bin/env python3
"""
Generador de Templates de Documentos para Pipeline E-Learning
Crea plantillas Word y Excel para el flujo de validación
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from openpyxl import Workbook
from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import os

# Colores Davivienda
DAVI_RED = RGBColor(218, 41, 28)
DAVI_RED_HEX = "DA291C"

OUTPUT_DIR = "/Users/federico/Desktop/ia-davivienda/templates/documentos"


def set_heading_style(doc):
    """Configura estilos de encabezado"""
    styles = doc.styles

    # Título principal
    title_style = styles['Title']
    title_style.font.color.rgb = DAVI_RED
    title_style.font.size = Pt(24)
    title_style.font.bold = True

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.color.rgb = DAVI_RED
    h1.font.size = Pt(14)
    h1.font.bold = True


def add_header_with_logo(doc, title):
    """Agrega encabezado con título"""
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("DAVIVIENDA")
    run.font.size = Pt(12)
    run.font.color.rgb = DAVI_RED
    run.font.bold = True

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = DAVI_RED
    title_run.font.bold = True

    # Línea separadora
    doc.add_paragraph("═" * 70)


def add_section_header(doc, title):
    """Agrega encabezado de sección"""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = DAVI_RED
    doc.add_paragraph("─" * 60)


def add_field(doc, label, width=40):
    """Agrega campo para llenar"""
    para = doc.add_paragraph()
    para.add_run(f"{label}: ").bold = True
    para.add_run("_" * width)


def add_checkbox_field(doc, label):
    """Agrega campo con checkbox"""
    para = doc.add_paragraph()
    para.add_run("[ ] ").font.size = Pt(11)
    para.add_run(label)


def add_checkbox_row(doc, label, options=["OK", "FALTA", "N/A"]):
    """Agrega fila con múltiples checkboxes"""
    para = doc.add_paragraph()
    para.add_run(f"{label}").font.size = Pt(10)
    para.add_run("  " * 3)
    for opt in options:
        para.add_run(f"[ ] {opt}  ")


# ============================================
# TEMPLATE 1: SOLICITUD
# ============================================
def create_solicitud_template():
    doc = Document()
    add_header_with_logo(doc, "SOLICITUD DE CURSO E-LEARNING")

    # Información General
    add_section_header(doc, "INFORMACIÓN GENERAL")
    add_field(doc, "Área solicitante")
    add_field(doc, "Responsable del área")
    add_field(doc, "Correo de contacto")
    add_field(doc, "Fecha de solicitud")
    add_field(doc, "Fecha esperada de entrega")

    doc.add_paragraph()

    # Datos del Curso
    add_section_header(doc, "DATOS DEL CURSO")
    add_field(doc, "Nombre del módulo")
    add_field(doc, "Audiencia objetivo (¿quién tomará el curso?)")

    para = doc.add_paragraph()
    para.add_run("Nivel: ").bold = True
    para.add_run("[ ] Básico    [ ] Intermedio    [ ] Avanzado")

    add_field(doc, "Duración estimada (minutos)")

    doc.add_paragraph()

    # Objetivos
    add_section_header(doc, "OBJETIVOS DE APRENDIZAJE")
    para = doc.add_paragraph()
    para.add_run("¿Qué debe lograr el participante al finalizar el curso?").italic = True
    doc.add_paragraph()

    add_field(doc, "Objetivo principal", 50)
    doc.add_paragraph()
    doc.add_paragraph("Objetivos específicos:")
    for i in range(1, 5):
        add_field(doc, f"  {i}", 50)

    doc.add_paragraph()

    # Contenido
    add_section_header(doc, "TEMAS Y CONTENIDO")
    para = doc.add_paragraph()
    para.add_run("Liste los temas principales que debe cubrir el curso:").italic = True
    doc.add_paragraph()

    for i in range(1, 7):
        add_field(doc, f"Tema {i}", 50)

    doc.add_paragraph()

    # Recursos
    add_section_header(doc, "RECURSOS Y MATERIAL DE APOYO")
    para = doc.add_paragraph()
    para.add_run("Marque los recursos que adjunta a esta solicitud:").italic = True
    doc.add_paragraph()

    add_checkbox_field(doc, "Documento con contenido técnico/conceptual")
    add_checkbox_field(doc, "Presentaciones PowerPoint existentes")
    add_checkbox_field(doc, "Videos o material audiovisual")
    add_checkbox_field(doc, "Manuales o guías de referencia")
    add_checkbox_field(doc, "Links a recursos en línea")
    add_checkbox_field(doc, "Experto temático disponible para consultas")

    doc.add_paragraph()
    add_field(doc, "Enlaces a carpeta compartida (Google Drive, SharePoint, etc.)", 30)

    doc.add_paragraph()

    # Consideraciones
    add_section_header(doc, "CONSIDERACIONES ADICIONALES")
    para = doc.add_paragraph()
    para.add_run("¿Hay algún requisito especial o consideración que debamos tener en cuenta?").italic = True
    doc.add_paragraph()
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)

    doc.add_paragraph()

    # Firmas
    add_section_header(doc, "FIRMAS")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Solicitante:"
    table.cell(0, 1).text = "Fecha:"
    table.cell(1, 0).text = "\n_______________________"
    table.cell(1, 1).text = "\n_______________________"

    doc.save(os.path.join(OUTPUT_DIR, "01_Template_Solicitud.docx"))
    print("✓ 01_Template_Solicitud.docx")


# ============================================
# TEMPLATE 2: CHECKLIST V1
# ============================================
def create_v1_checklist_template():
    doc = Document()
    add_header_with_logo(doc, "CHECKLIST DE VALIDACIÓN V1")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run("Validación de Solicitud - Diseño Instruccional").italic = True

    doc.add_paragraph()

    # Info del módulo
    add_section_header(doc, "INFORMACIÓN DEL MÓDULO")
    add_field(doc, "Nombre del módulo")
    add_field(doc, "Área solicitante")
    add_field(doc, "Validador (DI)")
    add_field(doc, "Fecha de validación")

    doc.add_paragraph()

    # Checklist
    add_section_header(doc, "REVISIÓN DE COMPLETITUD")

    items = [
        "Información del área está completa",
        "Objetivo principal es claro y medible",
        "Audiencia está bien definida",
        "Nivel del curso es apropiado",
        "Temas listados cubren el alcance",
        "Contenido bruto/material está adjunto",
        "Links y recursos funcionan correctamente",
        "Duración estimada es realista",
        "Hay experto disponible para consultas",
    ]

    table = doc.add_table(rows=len(items)+1, cols=4)
    table.style = 'Table Grid'

    # Header
    headers = ["Criterio", "OK", "FALTA", "N/A"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    # Items
    for i, item in enumerate(items, 1):
        table.cell(i, 0).text = item
        table.cell(i, 1).text = "[ ]"
        table.cell(i, 2).text = "[ ]"
        table.cell(i, 3).text = "[ ]"

    doc.add_paragraph()

    # Observaciones
    add_section_header(doc, "OBSERVACIONES PARA EL ÁREA")
    para = doc.add_paragraph()
    para.add_run("Si marcó 'FALTA' en algún criterio, detalle qué se necesita:").italic = True
    doc.add_paragraph()
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)

    doc.add_paragraph()

    # Decisión
    add_section_header(doc, "DECISIÓN")
    doc.add_paragraph()
    add_checkbox_field(doc, "✅ APROBADO - Proceder a generar Malla Curricular")
    add_checkbox_field(doc, "🔄 DEVOLVER AL ÁREA - Requiere completar información (ver observaciones)")

    doc.add_paragraph()
    doc.add_paragraph()

    # Firma
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Firma Diseño Instruccional:"
    table.cell(0, 1).text = "Fecha:"
    table.cell(1, 0).text = "\n_______________________"
    table.cell(1, 1).text = "\n_______________________"

    doc.save(os.path.join(OUTPUT_DIR, "01_Template_V1_Checklist.docx"))
    print("✓ 01_Template_V1_Checklist.docx")


# ============================================
# TEMPLATE 3: MALLA CURRICULAR (Excel)
# ============================================
def create_malla_template():
    wb = Workbook()
    ws = wb.active
    ws.title = "Malla Curricular"

    # Estilos
    header_fill = PatternFill(start_color=DAVI_RED_HEX, end_color=DAVI_RED_HEX, fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    validation_fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Headers
    headers = [
        "ID", "Etapa", "Bloque", "Objetivo de Aprendizaje",
        "Habilidad", "Tipo de Habilidad", "Nombre del Recurso",
        "Tipo de Recurso", "Descripción", "Duración (min)",
        "Evaluación", "Enlaces/Referencias",
        "COMENTARIOS DI", "COMENTARIOS ÁREA", "ESTADO"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border

    # Columnas de validación con color diferente
    for col in range(13, 16):
        ws.cell(row=1, column=col).fill = validation_fill
        ws.cell(row=1, column=col).font = Font(bold=True, size=11)

    # Datos de ejemplo
    sample_data = [
        ["1", "Introducción", "Bienvenida", "Conocer el propósito del módulo",
         "Contextualización", "Conocimiento del negocio", "Video de bienvenida",
         "Video avatar", "Avatar presenta el módulo", "1", "-", "",
         "", "", ""],
        ["2", "Introducción", "Objetivos", "Identificar qué aprenderá",
         "Comprensión", "Conocimiento técnico", "Pantalla interactiva",
         "Reveal buttons", "Objetivos y temas", "1", "-", "",
         "", "", ""],
        ["3", "Desarrollo", "Concepto 1", "Definir el concepto principal",
         "Conocimiento técnico", "Habilidad técnica", "Infografía",
         "Interactivo", "Explicación del concepto", "3", "Quiz 1", "",
         "", "", ""],
        ["4", "Desarrollo", "Concepto 2", "Aplicar el concepto en casos",
         "Aplicación práctica", "Habilidad técnica", "Caso práctico",
         "Escenario", "Caso de uso real", "3", "-", "",
         "", "", ""],
        ["5", "Cierre", "Resumen", "Sintetizar lo aprendido",
         "Síntesis", "Conocimiento del negocio", "Flashcards",
         "Interactivo", "Repaso de conceptos", "2", "-", "",
         "", "", ""],
        ["6", "Cierre", "Evaluación", "Demostrar comprensión",
         "Evaluación", "Conocimiento técnico", "Quiz final",
         "Quiz", "3 preguntas", "2", "Quiz final", "",
         "", "", ""],
    ]

    for row_num, row_data in enumerate(sample_data, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col_num, value=value)
            cell.border = border
            cell.alignment = Alignment(vertical='center', wrap_text=True)
            if col_num >= 13:
                cell.fill = PatternFill(start_color="FFFDE7", end_color="FFFDE7", fill_type="solid")

    # Ajustar anchos
    widths = [5, 12, 15, 30, 20, 20, 20, 15, 25, 10, 12, 20, 20, 20, 10]
    for i, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    # Altura de fila de header
    ws.row_dimensions[1].height = 40

    # Lista de tipos de habilidad (para referencia)
    ws2 = wb.create_sheet("Tipos de Habilidad")
    tipos = [
        ["Tipo de Habilidad", "Descripción"],
        ["Habilidades técnicas", "Conocimientos específicos del producto/servicio"],
        ["Habilidades blandas", "Comunicación, trabajo en equipo, liderazgo"],
        ["Conocimiento del negocio", "Contexto organizacional, procesos"],
        ["Conocimiento técnico", "Aspectos operativos y procedimentales"],
    ]
    for row_num, row_data in enumerate(tipos, 1):
        for col_num, value in enumerate(row_data, 1):
            cell = ws2.cell(row=row_num, column=col_num, value=value)
            if row_num == 1:
                cell.font = Font(bold=True)

    # Lista de tipos de recurso
    ws3 = wb.create_sheet("Tipos de Recurso")
    recursos = [
        ["Tipo de Recurso", "Descripción"],
        ["Video avatar", "Video con presentador virtual animado"],
        ["Interactivo", "Contenido con botones, reveal, hotspots"],
        ["Infografía", "Gráfico informativo visual"],
        ["Escenario", "Caso práctico con decisiones"],
        ["Quiz", "Preguntas de evaluación"],
        ["Flashcards", "Tarjetas de repaso"],
        ["Comparador", "Tabla comparativa interactiva"],
        ["Video", "Video tradicional sin avatar"],
        ["Documento", "PDF o recurso descargable"],
    ]
    for row_num, row_data in enumerate(recursos, 1):
        for col_num, value in enumerate(row_data, 1):
            cell = ws3.cell(row=row_num, column=col_num, value=value)
            if row_num == 1:
                cell.font = Font(bold=True)

    wb.save(os.path.join(OUTPUT_DIR, "02_Template_Malla_Curricular.xlsx"))
    print("✓ 02_Template_Malla_Curricular.xlsx")


# ============================================
# TEMPLATE 4: GUIÓN
# ============================================
def create_guion_template():
    doc = Document()
    add_header_with_logo(doc, "GUIÓN DEL CURSO")

    # Info general
    add_section_header(doc, "INFORMACIÓN GENERAL")
    add_field(doc, "Nombre del módulo")
    add_field(doc, "Versión")
    add_field(doc, "Fecha")
    add_field(doc, "Duración total estimada")
    add_field(doc, "Cantidad de cápsulas/secciones")

    doc.add_paragraph()
    doc.add_paragraph("═" * 70)
    doc.add_paragraph()

    # Cápsula ejemplo 1
    cap = doc.add_paragraph()
    cap.add_run("CÁPSULA 1: INTRODUCCIÓN").bold = True
    cap.runs[0].font.size = Pt(14)
    cap.runs[0].font.color.rgb = DAVI_RED

    doc.add_paragraph("─" * 60)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Duración:"
    table.cell(0, 1).text = "[__] segundos"
    table.cell(1, 0).text = "Tipo de recurso:"
    table.cell(1, 1).text = "[ ] Video avatar  [ ] Interactivo  [ ] Infografía  [ ] Otro: ____"
    table.cell(2, 0).text = "Objetivo:"
    table.cell(2, 1).text = "[Qué logrará el participante con esta cápsula]"
    table.cell(3, 0).text = "Evaluación:"
    table.cell(3, 1).text = "[ ] Sí  [ ] No    Tipo: ____________"

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.add_run("GUIÓN DETALLADO:").bold = True

    doc.add_paragraph()

    # Segmento 1
    seg = doc.add_paragraph()
    seg.add_run("[00:00 - 00:05]").bold = True
    doc.add_paragraph("Visual:  [Descripción de lo que se ve en pantalla]")
    doc.add_paragraph("Audio:   \"[Texto exacto que dice el narrador/avatar]\"")

    doc.add_paragraph()

    # Segmento 2
    seg = doc.add_paragraph()
    seg.add_run("[00:05 - 00:15]").bold = True
    doc.add_paragraph("Visual:  [Descripción de lo que se ve en pantalla]")
    doc.add_paragraph("Audio:   \"[Texto exacto que dice el narrador/avatar]\"")

    doc.add_paragraph()

    # Segmento 3
    seg = doc.add_paragraph()
    seg.add_run("[00:15 - 00:25]").bold = True
    doc.add_paragraph("Visual:  [Descripción de lo que se ve en pantalla]")
    doc.add_paragraph("Audio:   \"[Texto exacto que dice el narrador/avatar]\"")

    doc.add_paragraph()

    # Pregunta de evaluación
    sub = doc.add_paragraph()
    sub.add_run("PREGUNTA DE EVALUACIÓN (si aplica):").bold = True
    doc.add_paragraph()
    doc.add_paragraph("Pregunta: ________________________________________________")
    doc.add_paragraph("A) ________________________________________________")
    doc.add_paragraph("B) ________________________________________________")
    doc.add_paragraph("C) ________________________________________________")
    doc.add_paragraph("D) ________________________________________________")
    doc.add_paragraph("Respuesta correcta: [__]")
    doc.add_paragraph("Retroalimentación: ________________________________________________")

    doc.add_paragraph()
    doc.add_paragraph("═" * 70)
    doc.add_paragraph()

    # Cápsula ejemplo 2
    cap = doc.add_paragraph()
    cap.add_run("CÁPSULA 2: [NOMBRE]").bold = True
    cap.runs[0].font.size = Pt(14)
    cap.runs[0].font.color.rgb = DAVI_RED

    doc.add_paragraph("─" * 60)
    doc.add_paragraph("[Repetir estructura anterior para cada cápsula]")

    doc.add_paragraph()
    doc.add_paragraph("═" * 70)
    doc.add_paragraph()

    # Recursos adicionales
    add_section_header(doc, "RECURSOS ADICIONALES NECESARIOS")
    add_checkbox_field(doc, "Imágenes/iconos específicos: ________________________________")
    add_checkbox_field(doc, "Videos de apoyo: ________________________________")
    add_checkbox_field(doc, "Documentos descargables: ________________________________")
    add_checkbox_field(doc, "Links externos: ________________________________")

    doc.save(os.path.join(OUTPUT_DIR, "03_Template_Guion.docx"))
    print("✓ 03_Template_Guion.docx")


# ============================================
# TEMPLATE 5: VALIDACIÓN V3 (Guión)
# ============================================
def create_v3_revision_template():
    doc = Document()
    add_header_with_logo(doc, "REVISIÓN DE GUIÓN - VALIDACIÓN V3")

    # Info
    add_section_header(doc, "INFORMACIÓN")
    add_field(doc, "Nombre del módulo")
    add_field(doc, "Versión del guión")
    add_field(doc, "Validador DI")
    add_field(doc, "Validador Área")
    add_field(doc, "Fecha de revisión")

    doc.add_paragraph()

    # Revisión general
    add_section_header(doc, "REVISIÓN GENERAL")

    criterios = [
        "El guión cubre todos los objetivos de la malla",
        "La duración total es apropiada",
        "El tono es adecuado para la audiencia",
        "La secuencia de contenidos es lógica",
        "Las evaluaciones miden los objetivos",
    ]

    table = doc.add_table(rows=len(criterios)+1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Criterio"
    table.cell(0, 1).text = "OK"
    table.cell(0, 2).text = "Ajustar"
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, criterio in enumerate(criterios, 1):
        table.cell(i, 0).text = criterio
        table.cell(i, 1).text = "[ ]"
        table.cell(i, 2).text = "[ ]"

    doc.add_paragraph()

    # Revisión por cápsula
    add_section_header(doc, "REVISIÓN POR CÁPSULA")

    for i in range(1, 4):
        box = doc.add_paragraph()
        box.add_run(f"CÁPSULA {i}: [Nombre]").bold = True

        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        rows = [
            ("Precisión técnica del contenido", "[ ]", "[ ]"),
            ("Tono y lenguaje apropiados", "[ ]", "[ ]"),
            ("Duración de la cápsula", "[ ]", "[ ]"),
            ("Visuales descritos claramente", "[ ]", "[ ]"),
        ]

        for j, (criterio, ok, ajustar) in enumerate(rows):
            table.cell(j, 0).text = criterio
            table.cell(j, 1).text = ok
            table.cell(j, 2).text = ajustar

        doc.add_paragraph()
        doc.add_paragraph("Comentarios específicos:")
        doc.add_paragraph("_" * 65)
        doc.add_paragraph("_" * 65)
        doc.add_paragraph()

    # Ajustes requeridos
    add_section_header(doc, "LISTA DE AJUSTES REQUERIDOS")
    para = doc.add_paragraph()
    para.add_run("Detalle los cambios específicos que deben realizarse:").italic = True
    doc.add_paragraph()

    for i in range(1, 8):
        doc.add_paragraph(f"{i}. " + "_" * 60)

    doc.add_paragraph()

    # Decisión
    add_section_header(doc, "DECISIÓN")
    doc.add_paragraph()
    add_checkbox_field(doc, "✅ APROBADO - Proceder a Producción")
    add_checkbox_field(doc, "🔄 REQUIERE AJUSTES - Implementar cambios listados arriba")

    doc.add_paragraph()
    doc.add_paragraph()

    # Firmas
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Firma DI:"
    table.cell(0, 1).text = "Firma Área:"
    table.cell(0, 2).text = "Fecha:"
    table.cell(1, 0).text = "\n________________"
    table.cell(1, 1).text = "\n________________"
    table.cell(1, 2).text = "\n________________"

    doc.save(os.path.join(OUTPUT_DIR, "03_Template_V3_Revision_Guion.docx"))
    print("✓ 03_Template_V3_Revision_Guion.docx")


# ============================================
# TEMPLATE 6: VALIDACIÓN V4 (SCORM Final)
# ============================================
def create_v4_feedback_template():
    doc = Document()
    add_header_with_logo(doc, "FEEDBACK FINAL - VALIDACIÓN SCORM V4")

    # Info
    add_section_header(doc, "INFORMACIÓN DEL MÓDULO")
    add_field(doc, "Nombre del módulo")

    para = doc.add_paragraph()
    para.add_run("Versión SCORM: ").bold = True
    para.add_run("[ ] v1    [ ] v2    [ ] Final")

    add_field(doc, "Fecha de prueba")

    para = doc.add_paragraph()
    para.add_run("Probado en LMS Territorium: ").bold = True
    para.add_run("[ ] Sí    [ ] No")

    add_field(doc, "Probador DI")
    add_field(doc, "Probador Área")

    doc.add_paragraph()

    # Checklist técnico
    add_section_header(doc, "CHECKLIST TÉCNICO (Diseño Instruccional)")

    tecnicos = [
        "El SCORM carga correctamente en el LMS",
        "Los videos reproducen sin errores",
        "El audio se escucha claramente",
        "Los subtítulos están sincronizados",
        "Los elementos interactivos funcionan",
        "La navegación entre secciones es fluida",
        "El quiz registra las respuestas",
        "El progreso se guarda correctamente",
        "El certificado/resultado se muestra al final",
    ]

    table = doc.add_table(rows=len(tecnicos)+1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Criterio Técnico"
    table.cell(0, 1).text = "OK"
    table.cell(0, 2).text = "ERROR"
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, item in enumerate(tecnicos, 1):
        table.cell(i, 0).text = item
        table.cell(i, 1).text = "[ ]"
        table.cell(i, 2).text = "[ ]"

    doc.add_paragraph()

    # Checklist contenido
    add_section_header(doc, "REVISIÓN DE CONTENIDO (Área Solicitante)")

    contenido = [
        "La información técnica es correcta y actualizada",
        "Los ejemplos son apropiados y relevantes",
        "El lenguaje es adecuado para la audiencia",
        "Las preguntas del quiz son correctas",
        "Las respuestas correctas están bien marcadas",
        "No hay errores de ortografía o redacción",
        "Los recursos visuales son apropiados",
    ]

    table = doc.add_table(rows=len(contenido)+1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Criterio de Contenido"
    table.cell(0, 1).text = "OK"
    table.cell(0, 2).text = "AJUSTAR"
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i, item in enumerate(contenido, 1):
        table.cell(i, 0).text = item
        table.cell(i, 1).text = "[ ]"
        table.cell(i, 2).text = "[ ]"

    doc.add_paragraph()

    # Errores encontrados
    add_section_header(doc, "ERRORES O AJUSTES ENCONTRADOS")
    para = doc.add_paragraph()
    para.add_run("Detalle cada error o ajuste necesario con la ubicación exacta:").italic = True
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "#"
    table.cell(0, 1).text = "Ubicación (sección/minuto)"
    table.cell(0, 2).text = "Descripción del error/ajuste"
    table.cell(0, 3).text = "Prioridad"
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    for i in range(1, 6):
        table.cell(i, 0).text = str(i)
        table.cell(i, 3).text = "[ ]Alta [ ]Media [ ]Baja"

    doc.add_paragraph()

    # Decisión
    add_section_header(doc, "DECISIÓN FINAL")
    doc.add_paragraph()
    add_checkbox_field(doc, "✅ APROBADO PARA PRODUCCIÓN - Subir versión final a LMS")
    add_checkbox_field(doc, "🔄 REQUIERE NUEVA VERSIÓN - Corregir errores listados")

    doc.add_paragraph()

    # Comentarios finales
    doc.add_paragraph("Comentarios adicionales:")
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)
    doc.add_paragraph("_" * 70)

    doc.add_paragraph()
    doc.add_paragraph()

    # Firmas
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Firma DI:"
    table.cell(0, 1).text = "Firma Área:"
    table.cell(0, 2).text = "Fecha:"
    table.cell(1, 0).text = "\n________________"
    table.cell(1, 1).text = "\n________________"
    table.cell(1, 2).text = "\n________________"

    doc.save(os.path.join(OUTPUT_DIR, "04_Template_V4_Feedback_Final.docx"))
    print("✓ 04_Template_V4_Feedback_Final.docx")


# ============================================
# MAIN
# ============================================
def main():
    print("\n📄 Generando templates de documentos...\n")

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    create_solicitud_template()
    create_v1_checklist_template()
    create_malla_template()
    create_guion_template()
    create_v3_revision_template()
    create_v4_feedback_template()

    print(f"\n✅ Todos los templates creados en:\n   {OUTPUT_DIR}\n")


if __name__ == "__main__":
    main()
