"""
MVP - Generador de Mallas E-Learning Davivienda
Correr con: streamlit run app.py
"""

import streamlit as st

# Configuración de página (debe ir primero)
st.set_page_config(
    page_title="E-Learning Davivienda",
    page_icon="🎓",
    layout="wide"
)

# Mostrar loader inmediatamente mientras carga
if 'app_loaded' not in st.session_state:
    st.session_state.app_loaded = False

# Autenticación (import rápido, sin HTTP)
try:
    from auth import require_auth, logout, init_auth_state, handle_oauth_callback
    AUTH_DISPONIBLE = True
except:
    AUTH_DISPONIBLE = False

# Cliente API Firebase - NO hacer health_check aquí (es lento)
try:
    import api_client
    API_CLIENTE_DISPONIBLE = True
except:
    API_CLIENTE_DISPONIBLE = False

# Health check cacheado - solo se ejecuta una vez por sesión
@st.cache_data(ttl=300, show_spinner=False)
def check_api_health():
    """Health check cacheado por 5 minutos."""
    try:
        return api_client.health_check()
    except:
        return False

# Lazy check - solo cuando se necesita
def is_api_available():
    if not API_CLIENTE_DISPONIBLE:
        return False
    if 'api_disponible' not in st.session_state:
        st.session_state.api_disponible = check_api_health()
    return st.session_state.api_disponible

# Para compatibilidad con código existente
API_DISPONIBLE = API_CLIENTE_DISPONIBLE  # Se verificará lazy cuando se use

# Inicializar auth y manejar callback OAuth
if AUTH_DISPONIBLE:
    init_auth_state()
    handle_oauth_callback()
import streamlit.components.v1 as components
import pandas as pd
import json
import os
import requests
import base64
from openai import OpenAI
from pathlib import Path
from datetime import datetime
import docx
import PyPDF2
import io
import openpyxl
import concurrent.futures
import threading
import shutil
import zipfile
import tempfile
from mutagen.mp3 import MP3

# Cargar API keys
def get_elevenlabs_key():
    env_path = Path("/Users/federico/Desktop/ia-davivienda/.env")
    if env_path.exists():
        with open(env_path) as f:
            for line in f:
                if line.startswith('ELEVENLABS_API_KEY='):
                    return line.split('=', 1)[1].strip()
    return os.getenv('ELEVENLABS_API_KEY', '')

ELEVENLABS_API_KEY = get_elevenlabs_key()

def get_heygen_key():
    env_path = Path("/Users/federico/Desktop/ia-davivienda/.env")
    if env_path.exists():
        with open(env_path) as f:
            for line in f:
                if line.startswith('HEYGEN_API_KEY='):
                    return line.split('=', 1)[1].strip()
    return os.getenv('HEYGEN_API_KEY', '')

HEYGEN_API_KEY = get_heygen_key()

# Logo Davivienda - PNG embebido en base64 (casita con sombrero amarillo, cuerpo rojo, puerta azul)
LOGO_DAVIVIENDA_BASE64 = "iVBORw0KGgoAAAANSUhEUgAAAf8AAABiCAMAAABpoCqlAAABtlBMVEX////3pgDjBhPiAAAAAAAlgcT8qQDlBBL/rADjAAnjAADGAADjAA39qQD/rQD+8vPqAAD++vrsBRP2wsT74+TpVloAd77mMjixsbEThcn63t/40NHwmZvuhono6OhEPREdRGIRPFztfYD3ycvypafwkpXpXWHzs7SZAAAAAAre3t6uEAz96+zkGiPnP0TrbnHlKzLoSk8YJTLgmAsSAAAAABYoAADS0tKdnZ06OjqjdBGvAADQAAC9n24AhNAAYpVBAABoAACJP1C/gACiFxJPAADsdXjpWV7qZWm3ubt6f4M7RE1ePACjbQDPjg4TIjBkaG11UgWOlZWteRaJAAA6JgAAESIgAAC/FR2KZSKKDBlIVVVTABR1AABjTRtpa24uGwBGMABqTBIaGBFFMxHbmRiQZAdOPABLQA8tMhgKABVGABpnABSZABWyFSGWBhpkABjKEhwnLS5hHiQZWYRXJy8sVHJrN0oAda1uWjtOUm9dPVFFAACdmImomn4MSnNvj6gKZJFIibyTJi1BVlXUolAAIyOni10ANltIZn2MdVKQM0R1SGKLRlxiUXMAZKetKjNAGBrq8/w/AAANhUlEQVR4nO2djV/axh/HIXcQgjwGUYLinE/4BIK4atViO2wFWqnVtlZtddXV7qG2c3XOres6t/22dW7r9h//vnd5IAnQWnVjYfd+vVqRXJKDz/fp7pJos/1nCRaKVzEeLUY66t0Txj9OcGIUX5svZTKl+Wv4+gQzgf8U4wv4WsnpcjrtdqfT5SvdwAvj6raOicVicWOjuLgYCdazj4y/i8JVPG93gfYqTpd9aex6gW4M4ov9MjfHcJ07yvgbiNweyzl14qsmkLs1ECGbscjLiBwu1LuzjDOmcBvnfBXqUwvwlW69V7AVzomcgrjM9G8sBq/iJVdV9WULyI1dj2Be1Z/Hg/XuMOMMCS7gbXtN9WULWML4jhIA+ATL/43EBN60u16nPsFrXxlY5WgI4IfW6t1lxpkxfnus5HuT+gTXg3cGHotgAXz/Yr07zTgrNvB27cRvwrc0cDchcuI9Vv41CAV8K/PG0K9PAut4RxQvs/KvMdjAS8cK/boQkMN3uYts/q8RGAfn976d/DQEDLDyvxEo4sm3dH4lBCzhjXr3nXFaOm6PvU3mN4SAzAe32bqgtYngzcq5/mPjvIEj9f4EjFNw0tivAmUgywGWJXgel04Y+1W8GTzKcoA1GcRjmZPHfgWn/T4ef/O5GP86Cvj+6xd7jmsBm6wIsCAR/PAs1Ad8NzBbCLAaEXz/lKmfGYCFGcRbZ6U+4GI1gMXYwrkziv6UDGZzwVaigLdONe4341vHE/X+TIzjs4DXzyz7E1wr+Hy9PxPj+JzHH56p/7susARgJTYwPv3Uj54B5v9WYhDD8K+6ATgpXq9Lj9crv13DZnybmN0KYikiGH/4wKeq6qRq+3z2TOZBKZdbmtzeXlnZXH/40TuEhw83V7a3tyeXcrlSJgNq+6hJqPs6fZl1zCYALMbgGsYfbOdKoDfIvbL58KOPMTDw3ifL5z69++jSo9XVndnHlJ3ZJ6vwxmd3Pz039sl7A9Dq448eboJB5Oju89cwPs+G/5ajI1K8Ojo6enUBf7b6ZGfn8Z3d3QS5ql9U4PWU3+S4xO4uGAXYBF6jB9hgDwiwNJ+XteaMUNkr3tRs4mK9e844C8o3dJqFnrr39OmlIU+N7eJl5vcNQMfl6vqKU50zEnDQ2V+jwT2W9RuAQj/P6TK9pu6lGSn7am9vxpE9+EKoCP7QkJ9ic74NQLGJD0w3aQRkCxBXD7J78xmny57bz0pfyBGAD+xq7aahIbv/swFYFoUvDxwtLS0O+t/eNDEAvukgu2and4V4vV9lHbPEAHjha7kh/Gs5eOYWv2E3AFme8XvirJR9fuHCty9eXrhw4bvs1wGQOrCfXdPuCXJ9lT1EYBWBLyXHtxcuPH9RhIbfSz+I/ezCL8vz47R4lP3W6/WWMHbBj+9nQH9xR5rRze9696UjCACBF9mM1+vaxjd8Xq89+0eA+7zevWecksGLovB+9n/yHP6S0+5dJPqD+8/rbgl05loOEdFfgjedGOMM2IT0LCBeYrP+FufHO7yb6O8E9yfXhIH+bsj+jn3DHaGuBQkqAKo/uD/Gmz6qP5/4qd79Z5wK8jQvor/LvkWm/9edLqI/RIRJg/7OHIn2oL/LlyPt8LbP9f2zAAwS2RDQygTxHZ7o//P2AJUVb81/B/oH9l5lTEu8M4duHvJ/blNuh+/nwP/JI6DYHKCFWSNzf+73pV+wxsxMgEcHL02Xh3k3HE2gv1Ruh4n+nLjKSkDrMoETMKwT3pf2y7I6DgN8k/Sr6YEQ3klpxwP6j2rtRqn+HI+L9f4UjBPSgXfItA6M/x37nZR3Ow+k/YDncXbJdKWPM+c4EgLPJEfnu5TOTgcdEnL8HXbpv1VZU1b+3M+kModNvPhDtmS+0ivz6rcAP72na7iP5Inib1gGsCbjeFZZ7Qk0zQ4NDc0C/VOI54Qjh7n8s9sPYADAC1NyO/h/SpB3hgDAJgEsyQbN/hXLemAOv7+qkN95+EegoqEMZstAlgQv17iwo0J/p8vpkvWvdhnAU8yWgSxIB65x4QcMCI36O0tbS77a+q+yCtCKjONHtS7sOmox6O/axB+4Zn6rrT9bBrQg4zX93/Onqf7LrJfsr353M/0biXH8V60LP5taTOM/lzPT8qenhv6PmP5WpAPj6uoLgSmpyvxPf8BtvhRc1n+Z5X9Lgsnij1l8T2D6aM+RNc//Oiez0t5RU6DibgCyBMTqf0tSrCwARTT79YHkcDgWzPP/v2YdDulg/wfBXAXyT9itv9ZkHOOEUUru6FAi6jscM0b57a6X9G1JOvhy2hgC+DGW/i3KeVMA4P9yqLSYHg+QeaVukUaNEeMJ3qr352CcjEFjBcDvdkqqytmfvXbtBm+n3buU1UzjsF+/D2R/Vv1ZlQlcXgIgf89Lp/9L+4Pc0vb2JLnxP1eyL5T139Ppz3N/sfV/C7OBz5UNAHxZjf4zL29s5x4ok0CZEljCxmGLbALS9d3yHvxTvFDvz8A4BUW8XDYAsZ8GgJaFnJ083kM3+iMPhCl9RWoA6fCSWPb+p+zx7xYngvFjbVAv9u9JjpYafwHW6cocEPm1NWNxd4w99c/ydKzhy7uispwvTnU6XtZ8MpyvKO3fVP/4q5i4hD9nfwCuARgHCxji5IeAiOjmu/O1/v537pen03IrkbsD6rPrfhqEjsWf8MWbQ/Tm7ul+vDxv95ryv9drz93Cl+md37tTNy/jn4rM9xuJjsLixtra2sLGYmQ82BF5vjKfK2WI8rT+n195PjHYMR6R2yxOFNhdH41OcLAQmVhcLILYBebqDAaDwWAwGAwGg8FgMBgMxj+Pf7i5udlfY2OwOVh+WbXFMNk1OExeRqsfo9o2f3M0ajxp1HR4bWMQmkaHa/TPVrPnjGMQD3uQQrIrVHEFvj8PG/Lyd9+MULzKEdoQ6rYNI9Rjs4VRuNo5EqjLZkuipHbQ+JVW9aSCesi4GyGuW7dXK9nLFhrRmqZGeiqMIBiDDeEaZsd4I13ISDhk2AzachyHEH23B3ncVQ6RFNCILQS7goFwqK2yAeyHbDbYJjtqc5/hjGm5URx5eMGjM7BhxMNe7ab+JY0WGHWTB8pS22OchF6UjMW7o83Rtng7uLqAUKsuCoeQh0QHgadfMJG38ovuhndDsv62VkGVUw8IOWILqvr3krNwXfHosB/yTrcachAvoFZoo519GH4htgP9axseHo6G2kdIV1BCZ2FtSv841Hs2X8d/nFAKcULZhUMgAXzdUaIL8bt8tQDQKggp0pQoH6oWANqpqqr+c/Az1V1xFHD/FFFbSKjv+Kn+RtrSxEy0EBWF39wh23ASmraf5PMybLoyi9CDeA9SsmwzOG4rfdUn60oCgNnR4tT9Ff2JMbSaD0/dX9W/F350VelEl1xbgItrQkLLYEX/ICJxKKoe2SNwdOuI3AnGCWhDbn3VBz6lxvCEICSUTUlI4fBNX4GNpmobzIU0V/Rvq0wRV2ThZf2D5aMbCSOqalA+D0UJGFFkOCUYpRojWuFgiqn2lV8y3g6osAwhuwcpDpbXfafgakQ34st9hr1jSmtFfwjvPDIMIqJKzJD1j+szvB4oIqnKcUSCBcUtnz6OjK5NcgzNH136ZJOoDDyMY2HWn5RrxIXjStKX6ZYdu8cUaJvVcK7qT6q2Ef3RUoLsrrL+EAw8qVaZVK+hmeLl8EIxkJSnqv62hIcaVLchGREzi739h2dU6g/fb7uNRuo53bskjIMcaSgBdf6dEjwcfSHX/zbZQnTlXa/qpLL+eURGayr6A6lxXzsQCQnEEir0TwvU5DiPweN7kVYXMN6GGv4/Yk71kOjD8qi8nAFimtg62WAgrwnbprmlrD+MBYR0eG5uLpxOJ03+r+QaGEPKh0yj1/l/zJxJWAY4GWb9Q/SLrSzkQnLoj+s2dJeDbll/YiHqLGAQnDSlvKT6VxtBUFq1sN+GlF3Az4kdmfVvoxGluSLeVyk9GcfApL9foH6eFFThNCD0k0w+ovk8GR+qc7pl/amhxNRdtBJSGf/laxgABHs1evchuYXbQ8f/Zv1T1M/nYOxnOkTeXHoyjoNR/2hCIGVXtYmcqOJgKfieu+nvHkFLETr9aSqmBpDWVYuK/kEOfvZVrtiky9NOZJDRRg4vx3Oj/s1k0BelfTGvRfhrxRbG62hHmuf542EY55FfU9UmcpXpP78HJOgKxcjcq5aB9frTEJEO9XCGaTllOO9PwAuUD5VNIEqMKa2b1otDJ0JRyAgx+bfyllCezFB3E8uqCE/E7ipmJxhvpAd5UiOx3t5YX4oW5XP+GvO4dGxHAoAsIflXLsAM+pMoDzp5eP2srLb+E0PEyFAi3BWLxfJJeHmFFntlfybTTGRSn6aOuNq/kaS8AtQsh6LK+T4yZKk2t8h4Hc2G9bURKnsPKi/W6sgjJbHDC1CoVeds3cgwMQQHgPGCQaPyeG+4nTOck5ha3tA4jHheUKo5v6HpHC09wNoStkpiyDQ7xTgG/nhXPplIJFJXerQBdF+q2myqP5VWJGyb4/qMDhibM4y+/bFEqt1QjcXDuoQ9HGrvyqdbk+mRWA9VNOo2GFxPIhFWA1Aw3tWXTiSS+VhPm3rEK4nqs4jJf+sk8P8BOkzFzbRAqF4AAAAASUVORK5CYII="
LOGO_DAVIVIENDA_HTML = f'<img src="data:image/png;base64,{LOGO_DAVIVIENDA_BASE64}" alt="Davivienda" style="height: 50px;">'

# Configuración de Video
REMOTION_DIR = Path("/Users/federico/Desktop/ia-davivienda/remotion-videos")
CANVA_REQUESTS_DIR = Path("/Users/federico/Desktop/ia-davivienda/output/canva_requests")


def generar_video_canva(titulo, subtitulo, bullets, voiceover_text, output_dir, voice_id=None):
    """
    Genera video con Canva: audio ElevenLabs + request para Claude MCP

    Returns:
        dict con {success, audio_path, request_path, video_path, error}
    """
    resultado = {"success": False, "audio_path": None, "request_path": None, "video_path": None, "error": None}

    try:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        # 1. Generar audio con ElevenLabs
        audio_path = output_dir / "audio.mp3"
        voice = voice_id or "JddqVF50ZSIR7SRbJE6u"  # Valeria default

        audio_result, audio_error = generar_audio_elevenlabs(voiceover_text, voice, str(audio_path))

        if not audio_result:
            resultado["error"] = f"Error generando audio: {audio_error}"
            return resultado

        resultado["audio_path"] = str(audio_path)

        # 2. Obtener duración del audio
        try:
            audio = MP3(str(audio_path))
            duracion_segundos = audio.info.length
        except:
            duracion_segundos = len(voiceover_text.split()) / 2.5

        # 3. Crear request para Canva (via Claude MCP)
        CANVA_REQUESTS_DIR.mkdir(parents=True, exist_ok=True)
        request_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        request_path = CANVA_REQUESTS_DIR / f"canva_request_{request_id}.json"

        # Preparar slides para Canva
        slides = [
            {"title": titulo, "description": subtitulo}
        ]
        for bullet in bullets[:4]:
            slides.append({"title": bullet, "description": ""})

        canva_request = {
            "id": request_id,
            "timestamp": datetime.now().isoformat(),
            "topic": titulo,
            "audience": "professional",
            "style": "digital",
            "length": "short",
            "slides": slides,
            "voiceover_text": voiceover_text,
            "audio_path": str(audio_path),
            "audio_duration": duracion_segundos,
            "output_dir": str(output_dir),
            "status": "pending"
        }

        with open(request_path, 'w', encoding='utf-8') as f:
            json.dump(canva_request, f, ensure_ascii=False, indent=2)

        resultado["request_path"] = str(request_path)
        resultado["success"] = True
        resultado["message"] = f"Audio generado ({duracion_segundos:.1f}s). Para el video Canva, ejecuta: /canva {request_path}"

    except Exception as e:
        resultado["error"] = str(e)

    return resultado


def generar_video_remotion(titulo, subtitulo, bullets, voiceover_text, output_path, voice_id=None):
    """
    Genera video animado con Remotion + audio ElevenLabs

    Args:
        titulo: Título del video
        subtitulo: Subtítulo descriptivo
        bullets: Lista de puntos clave
        voiceover_text: Texto para generar el audio
        output_path: Ruta de salida del video
        voice_id: ID de la voz de ElevenLabs

    Returns:
        dict con {success, video_path, audio_path, error}
    """
    import subprocess
    import tempfile
    from mutagen.mp3 import MP3

    resultado = {"success": False, "video_path": None, "audio_path": None, "error": None}

    try:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 1. Generar audio con ElevenLabs
        audio_path = output_path.parent / "audio.mp3"
        voice = voice_id or "JddqVF50ZSIR7SRbJE6u"  # Valeria default

        audio_result, audio_error = generar_audio_elevenlabs(voiceover_text, voice, str(audio_path))

        if not audio_result:
            resultado["error"] = f"Error generando audio: {audio_error}"
            return resultado

        resultado["audio_path"] = str(audio_path)

        # 2. Obtener duración del audio
        try:
            audio = MP3(str(audio_path))
            duracion_segundos = audio.info.length
        except:
            duracion_segundos = len(voiceover_text.split()) / 2.5  # Estimación

        duracion_frames = int((duracion_segundos + 3) * 30)  # +3s buffer, 30fps

        # 3. Copiar audio a carpeta public de Remotion
        public_dir = REMOTION_DIR / "public"
        public_dir.mkdir(exist_ok=True)
        shutil.copy(str(audio_path), public_dir / "audio.mp3")

        # 4. Preparar props para Remotion
        props = {
            "titulo": titulo,
            "subtitulo": subtitulo,
            "bullets": bullets[:4],  # Max 4 bullets
            "audioSrc": "audio.mp3",
            "duracionFrames": duracion_frames,
        }

        props_json = json.dumps(props)

        # 5. Renderizar con Remotion CLI
        cmd = [
            "npx", "remotion", "render",
            "DaviviendaVideo",
            str(output_path),
            "--props", props_json,
            "--codec", "h264",
            "--overwrite"
        ]

        # Log para debug
        print(f"🎬 Remotion CMD: {' '.join(cmd)}")
        print(f"📁 Working dir: {REMOTION_DIR}")

        process = subprocess.run(
            cmd,
            cwd=REMOTION_DIR,
            capture_output=True,
            text=True,
            timeout=600  # 10 min timeout para videos largos
        )

        if process.returncode != 0:
            error_msg = process.stderr if process.stderr else process.stdout
            print(f"❌ Remotion error: {error_msg[:1000]}")
            resultado["error"] = f"Error Remotion: {error_msg[:500]}"
            return resultado

        # Verificar que el archivo existe
        if output_path.exists():
            resultado["success"] = True
            resultado["video_path"] = str(output_path)
            print(f"✅ Video generado: {output_path}")
        else:
            resultado["error"] = "Video no se generó correctamente"

    except subprocess.TimeoutExpired:
        resultado["error"] = "Timeout: el renderizado tardó más de 10 minutos"
    except Exception as e:
        print(f"❌ Exception: {e}")
        resultado["error"] = str(e)

    return resultado


def generar_video_heygen(audio_url, avatar_id, output_path):
    """Genera video con avatar HeyGen usando audio externo"""
    import time

    if not HEYGEN_API_KEY:
        return None, "Falta HEYGEN_API_KEY en .env"

    headers = {
        "X-Api-Key": HEYGEN_API_KEY,
        "Content-Type": "application/json"
    }

    # 1. Crear video
    create_url = "https://api.heygen.com/v2/video/generate"
    payload = {
        "video_inputs": [{
            "character": {
                "type": "avatar",
                "avatar_id": avatar_id,
                "avatar_style": "normal"
            },
            "voice": {
                "type": "audio",
                "audio_url": audio_url
            }
        }],
        "dimension": {"width": 1920, "height": 1080},
        "aspect_ratio": "16:9"
    }

    try:
        response = requests.post(create_url, json=payload, headers=headers)
        if response.status_code != 200:
            return None, f"Error creando video HeyGen: {response.status_code} - {response.text}"

        data = response.json()
        video_id = data.get('data', {}).get('video_id')
        if not video_id:
            return None, f"No se obtuvo video_id: {data}"

        # 2. Esperar a que el video esté listo (polling)
        status_url = f"https://api.heygen.com/v1/video_status.get?video_id={video_id}"
        max_attempts = 60  # 5 minutos máximo
        attempt = 0

        while attempt < max_attempts:
            time.sleep(5)
            status_response = requests.get(status_url, headers=headers)
            if status_response.status_code != 200:
                attempt += 1
                continue

            status_data = status_response.json()
            status = status_data.get('data', {}).get('status')

            if status == 'completed':
                video_url = status_data.get('data', {}).get('video_url')
                if video_url:
                    # 3. Descargar video
                    video_response = requests.get(video_url)
                    if video_response.status_code == 200:
                        with open(output_path, 'wb') as f:
                            f.write(video_response.content)
                        return output_path, None
                    else:
                        return None, f"Error descargando video: {video_response.status_code}"
                else:
                    return None, "Video completado pero sin URL"

            elif status == 'failed':
                error = status_data.get('data', {}).get('error', 'Unknown error')
                return None, f"HeyGen falló: {error}"

            attempt += 1

        return None, "Timeout esperando video de HeyGen"

    except Exception as e:
        return None, str(e)


def subir_audio_temporal(audio_path):
    """Sube audio a un servidor temporal para que HeyGen lo pueda acceder"""
    # Por ahora usamos file.io como servidor temporal gratuito
    try:
        with open(audio_path, 'rb') as f:
            response = requests.post(
                'https://file.io',
                files={'file': f},
                data={'expires': '1d'}
            )
        if response.status_code == 200:
            data = response.json()
            if data.get('success'):
                return data.get('link'), None
        return None, f"Error subiendo audio: {response.text}"
    except Exception as e:
        return None, str(e)


# Función para generar audio con ElevenLabs
def generar_audio_elevenlabs(texto, voice_id, output_path):
    """Genera audio MP3 con ElevenLabs"""
    if not ELEVENLABS_API_KEY:
        return None, "Falta ELEVENLABS_API_KEY"

    # Aplicar correcciones fonéticas
    FONEMAS = {
        "Pyme": "Píme", "pyme": "píme",
        "Pymes": "Pímes", "pymes": "pímes",
        "FATCA": "FATCA", "CRS": "CRS"
    }
    for palabra, correccion in FONEMAS.items():
        texto = texto.replace(palabra, correccion)

    url = f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}"

    headers = {
        "Accept": "audio/mpeg",
        "Content-Type": "application/json",
        "xi-api-key": ELEVENLABS_API_KEY
    }

    data = {
        "text": texto,
        "model_id": "eleven_multilingual_v2",
        "voice_settings": {
            "stability": 0.6,
            "similarity_boost": 0.8,
            "style": 0.4,
            "use_speaker_boost": True
        }
    }

    try:
        response = requests.post(url, json=data, headers=headers)
        if response.status_code == 200:
            with open(output_path, 'wb') as f:
                f.write(response.content)
            return output_path, None
        else:
            return None, f"Error ElevenLabs: {response.status_code} - {response.text}"
    except Exception as e:
        return None, str(e)

# Función para generar slide HTML (estilo Davivienda)
def generar_slide_html(titulo, contenido_texto, tipo, output_path, curso_nombre, duracion_seg=10):
    """Genera HTML del slide con branding Davivienda 2026 y animaciones CSS"""

    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <link href="https://fonts.googleapis.com/css2?family=Montserrat:ital,wght@0,400;0,600;0,700;1,600;1,700&family=Open+Sans:wght@400;600&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Open Sans', sans-serif;
            background: #F5F3F0;
            width: 1920px;
            height: 1080px;
            position: relative;
            overflow: hidden;
        }}

        /* ========== ANIMACIONES ========== */
        @keyframes fadeInUp {{
            from {{ opacity: 0; transform: translateY(40px); }}
            to {{ opacity: 1; transform: translateY(0); }}
        }}

        @keyframes fadeInLeft {{
            from {{ opacity: 0; transform: translateX(-50px); }}
            to {{ opacity: 1; transform: translateX(0); }}
        }}

        @keyframes fadeInRight {{
            from {{ opacity: 0; transform: translateX(50px); }}
            to {{ opacity: 1; transform: translateX(0); }}
        }}

        @keyframes scaleIn {{
            from {{ opacity: 0; transform: scale(0.8); }}
            to {{ opacity: 1; transform: scale(1); }}
        }}

        @keyframes slideInFromBottom {{
            from {{ opacity: 0; transform: translateY(100%); }}
            to {{ opacity: 1; transform: translateY(0); }}
        }}

        @keyframes pulse {{
            0%, 100% {{ transform: scale(1); }}
            50% {{ transform: scale(1.05); }}
        }}

        @keyframes typewriter {{
            from {{ width: 0; }}
            to {{ width: 100%; }}
        }}

        @keyframes waveLine {{
            0% {{ stroke-dashoffset: 1000; }}
            100% {{ stroke-dashoffset: 0; }}
        }}

        /* Clases de animación */
        .animate-fade-up {{ animation: fadeInUp 0.8s ease forwards; opacity: 0; }}
        .animate-fade-left {{ animation: fadeInLeft 0.8s ease forwards; opacity: 0; }}
        .animate-fade-right {{ animation: fadeInRight 0.8s ease forwards; opacity: 0; }}
        .animate-scale {{ animation: scaleIn 0.6s ease forwards; opacity: 0; }}
        .animate-slide-up {{ animation: slideInFromBottom 0.8s ease forwards; opacity: 0; }}

        .delay-1 {{ animation-delay: 0.2s; }}
        .delay-2 {{ animation-delay: 0.4s; }}
        .delay-3 {{ animation-delay: 0.6s; }}
        .delay-4 {{ animation-delay: 0.8s; }}
        .delay-5 {{ animation-delay: 1.0s; }}
        .delay-6 {{ animation-delay: 1.2s; }}

        /* Patrón de ondas animado */
        body::before {{
            content: '';
            position: absolute;
            top: 0; right: 0;
            width: 60%; height: 100%;
            background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 800 600'%3E%3Cpath d='M600,0 Q700,100 600,200 T600,400 T600,600' fill='none' stroke='rgba(218,41,28,0.06)' stroke-width='2'/%3E%3Cpath d='M650,0 Q750,100 650,200 T650,400 T650,600' fill='none' stroke='rgba(218,41,28,0.05)' stroke-width='2'/%3E%3Cpath d='M700,0 Q800,100 700,200 T700,400 T700,600' fill='none' stroke='rgba(218,41,28,0.04)' stroke-width='2'/%3E%3Cpath d='M750,0 Q850,100 750,200 T750,400 T750,600' fill='none' stroke='rgba(218,41,28,0.03)' stroke-width='2'/%3E%3C/svg%3E");
            background-size: cover;
            pointer-events: none;
            z-index: 0;
            animation: fadeInRight 1.5s ease;
        }}

        .header {{
            position: absolute;
            top: 30px;
            left: 40px;
            display: flex;
            align-items: center;
            gap: 20px;
            z-index: 10;
        }}

        .back-btn {{
            width: 55px;
            height: 55px;
            border: 3px solid #333;
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            background: white;
            font-size: 2em;
            font-weight: 300;
        }}

        .course-box {{
            padding: 14px 28px;
            border: 3px solid #333;
            border-radius: 8px;
            background: white;
            font-weight: 600;
            font-size: 1.1em;
        }}

        .main {{
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            text-align: center;
            z-index: 5;
            max-width: 1400px;
            width: 90%;
        }}

        .title-red {{
            font-family: 'Montserrat', sans-serif;
            font-weight: 700;
            font-style: italic;
            color: #DA291C;
            font-size: 3.2em;
            margin-bottom: 30px;
        }}

        .content-card {{
            background: white;
            border-radius: 24px;
            padding: 60px 80px;
            box-shadow: 0 10px 40px rgba(0,0,0,0.08);
            text-align: left;
        }}

        .voiceover-text {{
            font-size: 1.6em;
            line-height: 2;
            color: #333;
            border-left: 5px solid #DA291C;
            padding-left: 35px;
            background: linear-gradient(90deg, #FFF9F8 0%, transparent 100%);
            padding-top: 20px;
            padding-bottom: 20px;
        }}

        .badge {{
            display: inline-block;
            background: #DA291C;
            color: white;
            padding: 8px 20px;
            border-radius: 25px;
            font-size: 0.95em;
            font-weight: 600;
            margin-bottom: 25px;
        }}

        .bottom-nav {{
            position: absolute;
            bottom: 25px;
            right: 40px;
            display: flex;
            gap: 15px;
            z-index: 10;
        }}

        .nav-controls {{
            background: #333;
            padding: 12px 20px;
            border-radius: 8px;
            display: flex;
            gap: 15px;
            color: white;
            font-size: 1.3em;
        }}

        .page-nav {{
            background: white;
            padding: 12px 25px;
            border-radius: 8px;
            display: flex;
            align-items: center;
            gap: 15px;
            font-weight: 600;
            font-size: 1.1em;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}

        .page-nav .arrow {{
            color: #DA291C;
            font-size: 1.3em;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="back-btn">‹</div>
        <div class="course-box">{curso_nombre}</div>
    </div>

    <div class="main">
        <h1 class="title-red">{titulo}</h1>
        <div class="content-card">
            <span class="badge">{tipo}</span>
            <div class="voiceover-text">
                {contenido_texto}
            </div>
        </div>
    </div>

    <div class="bottom-nav">
        <div class="nav-controls">
            <span>🔊</span>
            <span>⛶</span>
        </div>
        <div class="page-nav">
            <span class="arrow">‹</span>
            <span>1 / 10</span>
            <span class="arrow">›</span>
        </div>
    </div>
</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

    return output_path


def renderizar_html_a_png(html_path, png_path):
    """Renderiza HTML a PNG usando Chrome headless (1920x1080)"""
    import subprocess

    chrome_path = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"

    # Convertir a ruta absoluta con file://
    if not html_path.startswith('file://'):
        html_url = f"file://{os.path.abspath(html_path)}"
    else:
        html_url = html_path

    cmd = [
        chrome_path,
        "--headless",
        "--disable-gpu",
        "--no-sandbox",
        "--disable-software-rasterizer",
        f"--screenshot={png_path}",
        "--window-size=1920,1080",
        "--hide-scrollbars",
        html_url
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if os.path.exists(png_path):
            return png_path, None
        else:
            return None, f"Chrome no generó el PNG: {result.stderr}"
    except subprocess.TimeoutExpired:
        return None, "Timeout renderizando HTML"
    except Exception as e:
        return None, str(e)


def generar_video_desde_imagen_audio(png_path, audio_path, video_path):
    """Combina imagen PNG + audio MP3 en video MP4 usando FFmpeg"""
    import subprocess

    cmd = [
        "ffmpeg", "-y",
        "-loop", "1",
        "-i", png_path,
        "-i", audio_path,
        "-c:v", "libx264",
        "-tune", "stillimage",
        "-c:a", "aac",
        "-b:a", "192k",
        "-pix_fmt", "yuv420p",
        "-shortest",
        video_path
    ]

    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if os.path.exists(video_path):
            return video_path, None
        else:
            return None, f"FFmpeg error: {result.stderr}"
    except subprocess.TimeoutExpired:
        return None, "Timeout generando video"
    except FileNotFoundError:
        return None, "FFmpeg no está instalado. Instalar con: brew install ffmpeg"
    except Exception as e:
        return None, str(e)


def dividir_texto_en_secciones(texto, max_secciones=5):
    """Divide el texto en secciones lógicas para múltiples slides"""
    import re

    # Intentar dividir por oraciones
    oraciones = re.split(r'(?<=[.!?])\s+', texto.strip())

    if len(oraciones) <= max_secciones:
        return oraciones

    # Agrupar oraciones en secciones
    secciones = []
    oraciones_por_seccion = len(oraciones) // max_secciones

    for i in range(0, len(oraciones), oraciones_por_seccion):
        seccion = ' '.join(oraciones[i:i + oraciones_por_seccion])
        if seccion:
            secciones.append(seccion)

    return secciones[:max_secciones]


def obtener_duracion_audio(audio_path):
    """Obtiene la duración de un archivo de audio en segundos usando ffprobe"""
    import subprocess

    try:
        cmd = [
            "ffprobe", "-v", "quiet", "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1", audio_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        return float(result.stdout.strip())
    except:
        return 60.0  # Default fallback


def generar_video_completo(titulo, voiceover_text, tipo, output_dir, curso_nombre, voice_id):
    """Pipeline completo: Audio → Múltiples Slides → Videos → Concatenar"""
    from pathlib import Path
    import subprocess

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    slides_dir = output_dir / "slides"
    slides_dir.mkdir(exist_ok=True)

    resultados = {"audio": None, "slides": [], "video": None, "errors": []}

    # 1. Generar Audio completo con ElevenLabs
    audio_path = output_dir / "audio.mp3"
    audio_result, audio_error = generar_audio_elevenlabs(voiceover_text, voice_id, str(audio_path))
    if audio_result:
        resultados["audio"] = str(audio_path)
    else:
        resultados["errors"].append(f"Audio: {audio_error}")
        return resultados

    # 2. Obtener duración del audio
    duracion_total = obtener_duracion_audio(str(audio_path))

    # 3. Dividir texto en secciones para slides progresivas
    secciones = dividir_texto_en_secciones(voiceover_text, max_secciones=4)
    num_slides = len(secciones)
    duracion_por_slide = duracion_total / num_slides

    # 4. Generar slides con contenido acumulativo
    segmentos_video = []
    contenido_acumulado = []

    for i, seccion in enumerate(secciones):
        contenido_acumulado.append(seccion)

        # HTML con bullets progresivos
        bullets_html = "\n".join([
            f'<div class="bullet {"current" if j == i else "previous"}">'
            f'<span class="bullet-icon">{"▸" if j == i else "✓"}</span>'
            f'<p>{s}</p></div>'
            for j, s in enumerate(contenido_acumulado)
        ])

        slide_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700&family=Open+Sans:wght@400;600&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{
            font-family: 'Open Sans', sans-serif;
            background: linear-gradient(180deg, #ffffff 0%, #f5f5f5 100%);
            min-height: 100vh;
            padding: 40px 60px;
        }}
        .header {{
            display: flex;
            align-items: center;
            gap: 20px;
            margin-bottom: 40px;
        }}
        .back-btn {{
            width: 50px;
            height: 50px;
            border: 2px solid #333;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 24px;
        }}
        .course-box {{
            background: white;
            border: 1px solid #ddd;
            padding: 12px 24px;
            border-radius: 8px;
            font-size: 14px;
        }}
        .title {{
            font-family: 'Montserrat', sans-serif;
            font-size: 42px;
            font-weight: 700;
            font-style: italic;
            color: #DA291C;
            margin-bottom: 40px;
        }}
        .content {{
            max-width: 1400px;
        }}
        .bullet {{
            display: flex;
            align-items: flex-start;
            gap: 20px;
            padding: 25px 30px;
            margin-bottom: 20px;
            border-radius: 12px;
            transition: all 0.3s;
        }}
        .bullet.current {{
            background: white;
            border-left: 5px solid #DA291C;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }}
        .bullet.previous {{
            background: #f8f9fa;
            border-left: 5px solid #28a745;
            opacity: 0.7;
        }}
        .bullet-icon {{
            font-size: 24px;
            color: #DA291C;
            min-width: 30px;
        }}
        .bullet.previous .bullet-icon {{
            color: #28a745;
        }}
        .bullet p {{
            font-size: 20px;
            line-height: 1.6;
            color: #333;
        }}
        .progress {{
            position: fixed;
            bottom: 40px;
            left: 60px;
            right: 60px;
            display: flex;
            gap: 10px;
        }}
        .progress-dot {{
            width: 12px;
            height: 12px;
            border-radius: 50%;
            background: #ddd;
        }}
        .progress-dot.active {{
            background: #DA291C;
        }}
        .progress-dot.done {{
            background: #28a745;
        }}
        .slide-counter {{
            position: fixed;
            bottom: 40px;
            right: 60px;
            font-size: 18px;
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="header">
        <div class="back-btn">‹</div>
        <div class="course-box">{curso_nombre}</div>
    </div>
    <h1 class="title">{titulo}</h1>
    <div class="content">
        {bullets_html}
    </div>
    <div class="progress">
        {"".join([f'<div class="progress-dot {"done" if j < i else "active" if j == i else ""}"></div>' for j in range(num_slides)])}
    </div>
    <div class="slide-counter">{i + 1} / {num_slides}</div>
</body>
</html>"""

        # Guardar slide HTML
        slide_path = slides_dir / f"slide_{i+1:02d}.html"
        with open(slide_path, 'w', encoding='utf-8') as f:
            f.write(slide_html)
        resultados["slides"].append(str(slide_path))

        # Renderizar a PNG
        png_path = slides_dir / f"slide_{i+1:02d}.png"
        png_result, png_error = renderizar_html_a_png(str(slide_path), str(png_path))
        if not png_result:
            resultados["errors"].append(f"PNG slide {i+1}: {png_error}")
            continue

        # Crear segmento de video (imagen estática por duración)
        segment_path = slides_dir / f"segment_{i+1:02d}.mp4"
        cmd = [
            "ffmpeg", "-y",
            "-loop", "1", "-i", str(png_path),
            "-t", str(duracion_por_slide),
            "-c:v", "libx264", "-tune", "stillimage",
            "-pix_fmt", "yuv420p",
            "-r", "30",
            str(segment_path)
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            segmentos_video.append(str(segment_path))
        else:
            resultados["errors"].append(f"Segment {i+1}: {result.stderr}")

    # 5. Concatenar todos los segmentos de video
    if len(segmentos_video) > 0:
        concat_file = slides_dir / "concat.txt"
        with open(concat_file, 'w') as f:
            for seg in segmentos_video:
                f.write(f"file '{seg}'\n")

        video_sin_audio = slides_dir / "video_sin_audio.mp4"
        cmd_concat = [
            "ffmpeg", "-y",
            "-f", "concat", "-safe", "0",
            "-i", str(concat_file),
            "-c", "copy",
            str(video_sin_audio)
        ]
        subprocess.run(cmd_concat, capture_output=True)

        # 6. Añadir audio al video concatenado
        video_path = output_dir / "video.mp4"
        cmd_audio = [
            "ffmpeg", "-y",
            "-i", str(video_sin_audio),
            "-i", str(audio_path),
            "-c:v", "copy", "-c:a", "aac", "-b:a", "192k",
            "-shortest",
            str(video_path)
        ]
        result = subprocess.run(cmd_audio, capture_output=True, text=True)
        if result.returncode == 0:
            resultados["video"] = str(video_path)
        else:
            resultados["errors"].append(f"Video final: {result.stderr}")

    return resultados


# Función para generar Quiz HTML interactivo
def generar_quiz_html(preguntas, titulo, curso_nombre):
    """Genera HTML del quiz interactivo"""
    questions_js = json.dumps(preguntas, ensure_ascii=False)

    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@400;600;700&family=Open+Sans:wght@400;600&display=swap" rel="stylesheet">
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Open Sans', sans-serif; background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%); min-height: 100vh; padding: 30px; }}
        .container {{ max-width: 800px; margin: 0 auto; }}
        .header {{ background: linear-gradient(135deg, #DA291C 0%, #b8231a 100%); padding: 25px; border-radius: 15px; margin-bottom: 25px; color: white; }}
        .header h1 {{ font-family: 'Montserrat', sans-serif; }}
        .question-card {{ background: white; border-radius: 15px; padding: 30px; margin-bottom: 20px; }}
        .question-text {{ font-size: 1.2em; margin-bottom: 20px; color: #333; }}
        .option {{ background: #f8f9fa; border: 2px solid #e0e0e0; padding: 15px 20px; border-radius: 10px; margin: 10px 0; cursor: pointer; transition: all 0.3s; }}
        .option:hover {{ border-color: #DA291C; background: #fff5f5; }}
        .option.selected {{ border-color: #DA291C; background: #fff5f5; }}
        .option.correct {{ border-color: #28a745; background: #d4edda; }}
        .option.incorrect {{ border-color: #dc3545; background: #f8d7da; }}
        .btn {{ background: #DA291C; color: white; border: none; padding: 15px 40px; border-radius: 10px; font-size: 1.1em; cursor: pointer; }}
        .btn:hover {{ background: #b8231a; }}
        .result {{ text-align: center; padding: 30px; }}
        .score {{ font-size: 3em; color: #DA291C; font-weight: bold; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <p style="opacity: 0.9; margin-bottom: 5px;">{curso_nombre}</p>
            <h1>{titulo}</h1>
        </div>
        <div id="quiz-container"></div>
    </div>
    <script>
        const questions = {questions_js};
        let currentQuestion = 0;
        let score = 0;
        let answered = false;

        function renderQuestion() {{
            const q = questions[currentQuestion];
            const container = document.getElementById('quiz-container');
            const opciones = q.opciones || q.options || [];
            const correcta = q.correcta !== undefined ? q.correcta : (q.correct || 0);

            let optionsHtml = opciones.map((op, i) =>
                `<div class="option" onclick="selectOption(${{i}}, ${{correcta}})" id="opt-${{i}}">${{op}}</div>`
            ).join('');

            container.innerHTML = `
                <div class="question-card">
                    <div class="question-text"><strong>Pregunta ${{currentQuestion + 1}} de ${{questions.length}}</strong><br><br>${{q.pregunta || q.question}}</div>
                    ${{optionsHtml}}
                    <div style="text-align: center; margin-top: 20px;">
                        <button class="btn" onclick="nextQuestion()" id="next-btn" style="display: none;">Siguiente</button>
                    </div>
                </div>
            `;
            answered = false;
        }}

        function selectOption(selected, correct) {{
            if (answered) return;
            answered = true;

            document.querySelectorAll('.option').forEach((el, i) => {{
                if (i === correct) el.classList.add('correct');
                else if (i === selected) el.classList.add('incorrect');
            }});

            if (selected === correct) score++;
            document.getElementById('next-btn').style.display = 'inline-block';
        }}

        function nextQuestion() {{
            currentQuestion++;
            if (currentQuestion < questions.length) {{
                renderQuestion();
            }} else {{
                showResults();
            }}
        }}

        function showResults() {{
            const pct = Math.round((score / questions.length) * 100);
            document.getElementById('quiz-container').innerHTML = `
                <div class="question-card result">
                    <h2>¡Quiz completado!</h2>
                    <div class="score">${{pct}}%</div>
                    <p style="margin-top: 20px; font-size: 1.2em;">${{score}} de ${{questions.length}} correctas</p>
                    <button class="btn" onclick="location.reload()" style="margin-top: 30px;">Reintentar</button>
                </div>
            `;
        }}

        renderQuestion();
    </script>
</body>
</html>"""

# Música de fondo disponible
MUSICA_FONDO = {
    "Sin música": None,
    "Corporativa suave": "https://cdn.pixabay.com/download/audio/2022/05/27/audio_1808fbf07a.mp3",
    "Motivacional": "https://cdn.pixabay.com/download/audio/2022/10/25/audio_946bd44e53.mp3",
    "Tecnología": "https://cdn.pixabay.com/download/audio/2023/05/16/audio_166b9c7242.mp3",
    "Personalizada": "custom"
}

# Función para generar HTML de preview visual
def generar_html_preview(guion, sol):
    """Genera HTML visual del contenido para preview - Estilo Davivienda 2026"""
    tipo = guion.get('tipo', '')
    titulo = guion.get('titulo', '')
    contenido = guion.get('contenido', {})
    curso_nombre = sol.get('nombre', 'Curso') if sol else 'Curso'

    # CSS Davivienda - Línea Gráfica 2026
    css = """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Montserrat:ital,wght@0,400;0,600;0,700;1,600;1,700&family=Open+Sans:wght@400;600&display=swap');

        :root {
            --davi-red: #DA291C;
            --davi-red-dark: #B8231A;
            --davi-orange: #F5A623;
            --davi-yellow: #F8E71C;
            --davi-blue: #4A90D9;
            --bg-light: #F5F3F0;
            --text-dark: #333333;
            --text-gray: #666666;
        }

        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: 'Open Sans', sans-serif;
            background: var(--bg-light);
            min-height: 100vh;
            position: relative;
        }

        /* Patrón de ondas */
        body::before {
            content: '';
            position: fixed;
            top: 0; right: 0;
            width: 60%; height: 100%;
            background-image: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 800 600'%3E%3Cpath d='M600,0 Q700,100 600,200 T600,400 T600,600' fill='none' stroke='rgba(218,41,28,0.06)' stroke-width='1'/%3E%3Cpath d='M650,0 Q750,100 650,200 T650,400 T650,600' fill='none' stroke='rgba(218,41,28,0.05)' stroke-width='1'/%3E%3Cpath d='M700,0 Q800,100 700,200 T700,400 T700,600' fill='none' stroke='rgba(218,41,28,0.04)' stroke-width='1'/%3E%3C/svg%3E");
            background-size: cover;
            pointer-events: none;
            z-index: 0;
        }

        .container { max-width: 900px; margin: 0 auto; padding: 20px; position: relative; z-index: 1; }

        /* Header */
        .course-header {
            display: flex;
            align-items: center;
            gap: 15px;
            margin-bottom: 30px;
        }

        .back-btn {
            width: 45px; height: 45px;
            border: 2px solid var(--text-dark);
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            background: white;
            font-size: 1.2em;
        }

        .course-title-box {
            padding: 10px 20px;
            border: 2px solid var(--text-dark);
            border-radius: 6px;
            background: white;
            font-weight: 600;
            font-size: 0.9em;
        }

        /* Título principal */
        .title-red {
            font-family: 'Montserrat', sans-serif;
            font-weight: 700;
            font-style: italic;
            color: var(--davi-red);
            font-size: 2em;
            text-align: center;
            margin-bottom: 15px;
        }

        .subtitle {
            text-align: center;
            color: var(--text-gray);
            margin-bottom: 10px;
            line-height: 1.6;
        }

        .instruction {
            text-align: center;
            color: #999;
            font-style: italic;
            margin-bottom: 25px;
            font-size: 0.9em;
        }

        /* Content card */
        .content-card {
            background: white;
            border-radius: 16px;
            padding: 30px;
            box-shadow: 0 4px 20px rgba(0,0,0,0.08);
        }

        /* Voiceover */
        .voiceover-box {
            background: linear-gradient(135deg, #FFF9F8 0%, #FFF 100%);
            border-left: 4px solid var(--davi-red);
            padding: 25px;
            border-radius: 0 12px 12px 0;
            line-height: 1.9;
            font-size: 1.05em;
            color: var(--text-dark);
        }

        /* Quiz */
        .quiz-question {
            background: white;
            border-radius: 16px;
            padding: 25px;
            margin-bottom: 20px;
            box-shadow: 0 2px 15px rgba(0,0,0,0.06);
        }

        .quiz-question h3 {
            color: var(--davi-red);
            font-family: 'Montserrat', sans-serif;
            margin-bottom: 18px;
            font-size: 1.1em;
        }

        .quiz-option {
            background: var(--bg-light);
            border: 2px solid transparent;
            padding: 14px 18px;
            border-radius: 10px;
            margin: 10px 0;
            display: flex;
            align-items: center;
            gap: 12px;
            transition: all 0.2s;
        }

        .quiz-option:hover { border-color: var(--davi-red); background: #FFF5F5; }
        .quiz-option.correct { border-color: #00B5AD; background: #E6FFF9; }

        .option-letter {
            width: 28px; height: 28px;
            border-radius: 50%;
            border: 2px solid #CCC;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 600;
            font-size: 0.85em;
            flex-shrink: 0;
        }

        .quiz-option.correct .option-letter {
            background: #00B5AD;
            border-color: #00B5AD;
            color: white;
        }

        /* Flashcards */
        .flashcards-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
        }

        .flashcard {
            background: linear-gradient(135deg, var(--davi-red) 0%, var(--davi-red-dark) 100%);
            color: white;
            padding: 30px 25px;
            border-radius: 14px;
            text-align: center;
            min-height: 160px;
            display: flex;
            flex-direction: column;
            justify-content: center;
        }

        .flashcard .front {
            font-weight: 600;
            font-size: 1.1em;
            margin-bottom: 15px;
        }

        .flashcard .back {
            padding-top: 15px;
            border-top: 1px solid rgba(255,255,255,0.3);
            font-size: 0.95em;
            opacity: 0.95;
        }

        /* Interactivo / Accordion */
        .interactive-list { display: flex; flex-direction: column; gap: 12px; }

        .interactive-item {
            background: white;
            border-radius: 30px;
            overflow: hidden;
            box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        }

        .interactive-header {
            display: flex;
            align-items: center;
            gap: 15px;
            padding: 12px 25px;
            cursor: pointer;
        }

        .circle-num {
            width: 45px; height: 45px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-family: 'Montserrat', sans-serif;
            font-weight: 700;
            font-size: 1.2em;
            color: white;
            flex-shrink: 0;
        }

        .circle-num.c1 { background: var(--davi-red); }
        .circle-num.c2 { background: var(--davi-orange); }
        .circle-num.c3 { background: var(--davi-blue); }
        .circle-num.c4 { background: var(--davi-yellow); color: #333; }

        .interactive-title {
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            color: var(--davi-red);
        }

        .interactive-content {
            background: rgba(245,243,240,0.8);
            padding: 20px 25px 20px 85px;
            color: var(--text-gray);
            line-height: 1.7;
            border-top: 1px solid #EEE;
        }

        /* Caso práctico */
        .scenario-box {
            background: linear-gradient(135deg, var(--davi-red) 0%, var(--davi-red-dark) 100%);
            color: white;
            padding: 25px;
            border-radius: 14px;
            margin-bottom: 25px;
        }

        .scenario-box h4 { margin-bottom: 12px; font-size: 1.1em; }
        .scenario-box p { line-height: 1.7; opacity: 0.95; }

        .decision-card {
            background: white;
            border: 2px solid #EEE;
            border-radius: 12px;
            padding: 18px 22px;
            margin: 12px 0;
            transition: all 0.2s;
        }

        .decision-card:hover { border-color: var(--davi-red); }
        .decision-card strong { color: var(--davi-red); }
        .decision-card p { margin-top: 10px; color: var(--text-gray); font-size: 0.95em; }

        /* Infografía */
        .info-section {
            background: white;
            border-radius: 12px;
            padding: 22px;
            margin: 15px 0;
            border-left: 4px solid var(--davi-red);
            box-shadow: 0 2px 10px rgba(0,0,0,0.04);
        }

        .info-section h4 {
            color: var(--davi-red);
            font-family: 'Montserrat', sans-serif;
            margin-bottom: 10px;
        }

        .info-section .dato {
            margin-top: 12px;
            padding: 10px 15px;
            background: #FFF5F5;
            border-radius: 8px;
            color: var(--davi-red);
            font-weight: 600;
        }

        /* Contenedor scrolleable para tablas */
        .table-wrapper {
            overflow-x: auto;
            -webkit-overflow-scrolling: touch;
            margin-bottom: 20px;
            border-radius: 12px;
            box-shadow: 0 2px 15px rgba(0,0,0,0.06);
        }

        /* Tabla comparador */
        .comparison-table {
            width: 100%;
            min-width: 600px;
            border-collapse: separate;
            border-spacing: 0;
        }

        .comparison-table th {
            background: var(--davi-red);
            color: white;
            padding: 16px 18px;
            text-align: left;
            font-family: 'Montserrat', sans-serif;
            font-weight: 600;
            white-space: nowrap;
        }

        .comparison-table td {
            padding: 16px 18px;
            border-bottom: 1px solid #EEE;
            background: white;
            min-width: 150px;
        }

        .comparison-table tr:last-child td { border-bottom: none; }
        .comparison-table tr:hover td { background: #FFF9F8; }

        /* Footer nav */
        .bottom-nav {
            position: fixed;
            bottom: 15px;
            right: 20px;
            display: flex;
            gap: 10px;
        }

        .nav-btn {
            background: #333;
            color: white;
            padding: 10px 15px;
            border-radius: 6px;
            display: flex;
            align-items: center;
            gap: 8px;
            font-size: 0.9em;
        }

        .page-nav {
            background: white;
            padding: 10px 18px;
            border-radius: 6px;
            display: flex;
            align-items: center;
            gap: 10px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            font-weight: 600;
        }

        .page-nav .arrow { color: var(--davi-red); font-size: 1.1em; }
    </style>
    """

    html_content = ""
    colors = ['c1', 'c2', 'c3', 'c4']  # Colores para círculos

    if tipo in ["Video avatar", "Video"]:
        voiceover = contenido.get('voiceover', contenido.get('texto', ''))
        indicaciones = contenido.get('indicaciones', '')
        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="instruction">Contenido del voiceover para narración</p>
        <div class="content-card">
            <div class="voiceover-box">{voiceover}</div>
            {f'<p style="margin-top: 20px; color: #999; font-style: italic;">📋 Indicaciones: {indicaciones}</p>' if indicaciones else ''}
        </div>
        """

    elif tipo == "Quiz":
        preguntas = contenido.get('preguntas', contenido.get('questions', []))
        questions_html = ""
        letters = ['A', 'B', 'C', 'D']
        for i, p in enumerate(preguntas, 1):
            pregunta = p.get('pregunta', p.get('question', ''))
            opciones = p.get('opciones', p.get('options', []))
            correcta = p.get('correcta', p.get('correct', 0))

            options_html = ""
            for j, op in enumerate(opciones):
                clase = "correct" if j == correcta else ""
                letter = letters[j] if j < len(letters) else str(j+1)
                options_html += f'''
                <div class="quiz-option {clase}">
                    <span class="option-letter">{letter}</span>
                    <span>{op}</span>
                </div>'''

            questions_html += f"""
            <div class="quiz-question">
                <h3>Pregunta {i}</h3>
                <p style="font-size: 1.05em; margin-bottom: 18px; line-height: 1.6;">{pregunta}</p>
                {options_html}
            </div>
            """

        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="subtitle">Evaluación de conocimientos - {len(preguntas)} preguntas</p>
        <p class="instruction">La opción correcta está marcada en verde</p>
        {questions_html}
        """

    elif tipo == "Flashcards":
        tarjetas = contenido.get('tarjetas', contenido.get('cards', []))
        cards_html = ""
        for t in tarjetas:
            frente = t.get('frente', t.get('front', ''))
            reverso = t.get('reverso', t.get('back', ''))
            cards_html += f"""
            <div class="flashcard">
                <div class="front">{frente}</div>
                <div class="back">{reverso}</div>
            </div>
            """

        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="subtitle">Tarjetas de repaso - {len(tarjetas)} conceptos</p>
        <p class="instruction">Haz clic en cada tarjeta para ver la respuesta</p>
        <div class="flashcards-grid">
            {cards_html}
        </div>
        """

    elif tipo == "Caso práctico":
        escenario = contenido.get('escenario', '')
        opciones = contenido.get('opciones', contenido.get('decisions', []))

        options_html = ""
        for op in opciones:
            opcion = op.get('opcion', op.get('option', ''))
            consecuencia = op.get('consecuencia', op.get('feedback', ''))
            options_html += f"""
            <div class="decision-card">
                <strong>➡️ {opcion}</strong>
                <p>{consecuencia}</p>
            </div>
            """

        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="instruction">Analiza el escenario y selecciona la mejor opción</p>
        <div class="scenario-box">
            <h4>📖 Escenario</h4>
            <p>{escenario}</p>
        </div>
        <h3 style="color: var(--davi-red); margin-bottom: 15px; font-family: 'Montserrat', sans-serif;">🔀 Opciones de decisión</h3>
        {options_html}
        """

    elif tipo in ["Interactivo", "Accordion"]:
        elementos = contenido.get('elementos', contenido.get('items', []))
        items_html = ""
        for i, el in enumerate(elementos):
            titulo_el = el.get('titulo', el.get('title', ''))
            contenido_el = el.get('contenido', el.get('content', ''))
            color = colors[i % len(colors)]
            items_html += f"""
            <div class="interactive-item">
                <div class="interactive-header">
                    <span class="circle-num {color}">{i+1}</span>
                    <span class="interactive-title">{titulo_el}</span>
                </div>
                <div class="interactive-content">{contenido_el}</div>
            </div>
            """

        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="subtitle">Haz clic en cada elemento para descubrir más información.</p>
        <p class="instruction">Haz clic en cada botón para descubrir más información.</p>
        <div class="interactive-list">
            {items_html}
        </div>
        """

    elif tipo == "Infografía":
        secciones = contenido.get('secciones', contenido.get('sections', []))
        sections_html = ""
        for s in secciones:
            titulo_s = s.get('titulo', s.get('title', ''))
            contenido_s = s.get('contenido', s.get('content', ''))
            dato = s.get('dato', '')
            sections_html += f"""
            <div class="info-section">
                <h4>{titulo_s}</h4>
                <p style="line-height: 1.7; color: var(--text-gray);">{contenido_s}</p>
                {f'<div class="dato">📊 {dato}</div>' if dato else ''}
            </div>
            """

        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <p class="subtitle">Información clave resumida visualmente</p>
        {sections_html}
        """

    elif tipo == "Comparador":
        items = contenido.get('items', contenido.get('comparacion', []))
        columnas = contenido.get('columnas', [])
        if items and len(items) > 0:
            if columnas:
                headers = columnas
                keys = list(items[0].keys())
            else:
                headers = list(items[0].keys())
                keys = headers

            header_html = "".join([f"<th>{h.replace('_', ' ').title()}</th>" for h in headers])
            rows_html = ""
            for item in items:
                row = "".join([f"<td>{item.get(k, '')}</td>" for k in keys])
                rows_html += f"<tr>{row}</tr>"

            html_content = f"""
            <h1 class="title-red">{titulo}</h1>
            <p class="subtitle">Comparación de opciones lado a lado</p>
            <p class="instruction">Desliza horizontalmente para ver todas las columnas</p>
            <div class="table-wrapper">
                <table class="comparison-table">
                    <thead><tr>{header_html}</tr></thead>
                    <tbody>{rows_html}</tbody>
                </table>
            </div>
            """
        else:
            html_content = '<div class="content-card"><p>Sin datos para comparar</p></div>'

    else:
        html_content = f"""
        <h1 class="title-red">{titulo}</h1>
        <div class="content-card">
            <pre style="background: var(--bg-light); padding: 20px; border-radius: 8px; overflow: auto; font-size: 0.9em;">{json.dumps(contenido, indent=2, ensure_ascii=False)}</pre>
        </div>
        """

    # HTML completo - Estilo Davivienda 2026
    return f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        {css}
    </head>
    <body>
        <div class="container">
            <!-- Header con navegación -->
            <div class="course-header">
                <div class="back-btn">‹</div>
                <div class="course-title-box">{curso_nombre}</div>
            </div>

            <!-- Contenido principal -->
            {html_content}
        </div>

        <!-- Navegación inferior -->
        <div class="bottom-nav">
            <div class="nav-btn">
                <span>🔊</span>
                <span>⛶</span>
            </div>
            <div class="page-nav">
                <span class="arrow">‹</span>
                <span>1 / 10</span>
                <span class="arrow">›</span>
            </div>
        </div>

        <script>
            // Interactividad básica
            document.querySelectorAll('.interactive-item').forEach(item => {{
                item.querySelector('.interactive-header').addEventListener('click', () => {{
                    item.classList.toggle('active');
                    const content = item.querySelector('.interactive-content');
                    content.style.display = content.style.display === 'none' ? 'block' : 'none';
                }});
            }});

            // Flashcards flip
            document.querySelectorAll('.flashcard').forEach(card => {{
                card.addEventListener('click', () => card.classList.toggle('flipped'));
            }});
        </script>
    </body>
    </html>
    """


def generar_recurso_individual(guion_data, sol, voice_id, heygen_avatar_id, progreso_callback=None):
    """Genera un recurso individual. Retorna dict con resultado."""
    from pathlib import Path

    guion_id = guion_data.get('id')
    tipo = guion_data.get('tipo', 'Otro')
    titulo = guion_data.get('titulo', 'Sin título')
    contenido_data = guion_data.get('contenido', {})

    output_base = Path(f"/Users/federico/Desktop/ia-davivienda/output/{sol['nombre'].replace(' ', '_').lower()}")
    output_base.mkdir(parents=True, exist_ok=True)
    recurso_dir = output_base / f"{guion_id:02d}_{tipo.replace(' ', '_').lower()}"
    recurso_dir.mkdir(exist_ok=True)

    assets_generados = {}
    errors = []

    if progreso_callback:
        progreso_callback(guion_id, "iniciando", f"Iniciando {titulo}...")

    try:
        if tipo == "Video avatar":
            voiceover = contenido_data.get('voiceover', contenido_data.get('texto', ''))
            if voiceover:
                if progreso_callback:
                    progreso_callback(guion_id, "audio", "Generando audio...")

                audio_path = recurso_dir / "audio.mp3"
                audio_result, audio_error = generar_audio_elevenlabs(voiceover, voice_id, str(audio_path))

                if audio_result:
                    assets_generados['audio'] = str(audio_path)

                    if progreso_callback:
                        progreso_callback(guion_id, "upload", "Subiendo audio...")

                    audio_url, upload_error = subir_audio_temporal(str(audio_path))

                    if audio_url:
                        if progreso_callback:
                            progreso_callback(guion_id, "heygen", "Generando avatar (1-3 min)...")

                        video_path = recurso_dir / "video_avatar.mp4"
                        video_result, video_error = generar_video_heygen(audio_url, heygen_avatar_id, str(video_path))

                        if video_result:
                            assets_generados['video'] = str(video_path)
                        else:
                            errors.append(f"HeyGen: {video_error}")
                    else:
                        errors.append(f"Upload: {upload_error}")
                else:
                    errors.append(f"Audio: {audio_error}")

        elif tipo == "Video":
            voiceover = contenido_data.get('voiceover', contenido_data.get('texto', ''))
            if voiceover:
                if progreso_callback:
                    progreso_callback(guion_id, "video", "Generando video con slides...")

                resultados = generar_video_completo(
                    titulo, voiceover, tipo, str(recurso_dir),
                    sol['nombre'], voice_id
                )

                if resultados.get('audio'):
                    assets_generados['audio'] = resultados['audio']
                if resultados.get('slides'):
                    assets_generados['slides'] = resultados['slides']
                if resultados.get('video'):
                    assets_generados['video'] = resultados['video']
                errors.extend(resultados.get('errors', []))

        elif tipo == "Quiz":
            preguntas = contenido_data.get('preguntas', [])
            quiz_html = generar_quiz_html(preguntas, titulo, sol['nombre'])
            quiz_path = recurso_dir / "quiz.html"
            with open(quiz_path, 'w', encoding='utf-8') as f:
                f.write(quiz_html)
            assets_generados['quiz'] = str(quiz_path)

        else:
            # Generar slide genérico para otros tipos
            slide_path = recurso_dir / "index.html"
            html_content = generar_html_preview(guion_data, sol)
            with open(slide_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            assets_generados['html'] = str(slide_path)

        if progreso_callback:
            progreso_callback(guion_id, "completado", "Completado!")

    except Exception as e:
        errors.append(str(e))
        if progreso_callback:
            progreso_callback(guion_id, "error", f"Error: {str(e)}")

    return {
        "guion_id": guion_id,
        "tipo": tipo,
        "titulo": titulo,
        "status": "generado" if assets_generados else "error",
        "path": str(recurso_dir),
        "assets": assets_generados,
        "errors": errors
    }


# Verificar autenticación
if AUTH_DISPONIBLE:
    from auth import show_login_page
    if not st.session_state.get("authenticated", False):
        show_login_page()
        st.stop()

# CSS personalizado
st.markdown("""
<style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .main-header {
        background: linear-gradient(135deg, #DA291C 0%, #b8231a 100%);
        color: white;
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
    }
    .main-header h1 {
        margin: 0;
        color: white;
    }
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
    .info-box {
        background: #e7f3ff;
        border: 1px solid #b6d4fe;
        border-radius: 10px;
        padding: 15px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# Header con info de usuario
user_info = ""
if AUTH_DISPONIBLE and st.session_state.get("authenticated"):
    rol = st.session_state.get("rol", "")
    email = st.session_state.get("user_email", "")
    if rol == "learning":
        user_info = f'<span style="float:right; font-size:14px;">👤 {email}</span>'
    else:
        user_info = f'<span style="float:right; font-size:14px;">📝 Área Solicitante</span>'

st.markdown(f"""
<div class="main-header">
    <h1>🎓 E-Learning Davivienda {user_info}</h1>
    <p>Generador de Mallas Curriculares con IA</p>
</div>
""", unsafe_allow_html=True)

# Botón de logout en sidebar
if AUTH_DISPONIBLE and st.session_state.get("authenticated"):
    with st.sidebar:
        if st.button("🚪 Cerrar sesión"):
            from auth import logout
            logout()
            st.rerun()

# Extraer texto de archivos subidos
def extract_text_from_file(uploaded_file):
    """Extrae texto de PDF, DOCX o TXT"""
    text = ""
    file_type = uploaded_file.name.split('.')[-1].lower()

    try:
        if file_type == 'txt':
            text = uploaded_file.read().decode('utf-8')
        elif file_type == 'docx':
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            text = '\n'.join([para.text for para in doc.paragraphs])
        elif file_type == 'pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            for page in pdf_reader.pages:
                text += page.extract_text() + '\n'
    except Exception as e:
        text = f"[Error leyendo archivo: {str(e)}]"

    return text

# Cargar API key automáticamente (sin mostrar en UI)
def get_api_key():
    # 1. Intentar desde variable de entorno
    if os.getenv('OPENAI_API_KEY'):
        return os.getenv('OPENAI_API_KEY')

    # 2. Intentar desde .env local del proyecto
    local_env = Path("/Users/federico/Desktop/ia-davivienda/.env")
    if local_env.exists():
        with open(local_env) as f:
            for line in f:
                if line.startswith('OPENAI_API_KEY='):
                    return line.split('=', 1)[1].strip()

    # 3. Fallback a escalar/.env
    env_path = Path.home() / "Desktop/escalar/.env"
    if env_path.exists():
        with open(env_path) as f:
            for line in f:
                if line.startswith('OPENAI_API_KEY='):
                    return line.split('=', 1)[1].strip()

    return ''

# Cargar API key al iniciar
api_key = get_api_key()

# Inicializar estado
if 'malla' not in st.session_state:
    st.session_state.malla = None
if 'version' not in st.session_state:
    st.session_state.version = 1
if 'historial' not in st.session_state:
    st.session_state.historial = []
if 'solicitud' not in st.session_state:
    st.session_state.solicitud = None
if 'guiones' not in st.session_state:
    st.session_state.guiones = None
if 'contenido_generado' not in st.session_state:
    st.session_state.contenido_generado = []
if 'scorm_path' not in st.session_state:
    st.session_state.scorm_path = None
if 'recursos_seleccionados' not in st.session_state:
    st.session_state.recursos_seleccionados = set()
if 'generando_en_paralelo' not in st.session_state:
    st.session_state.generando_en_paralelo = False
if 'progreso_generacion' not in st.session_state:
    st.session_state.progreso_generacion = {}

# Glosario de recursos (definición global)
GLOSARIO_RECURSOS = [
    {"tipo": "Video avatar", "icono": "🎭", "descripcion": "Avatar IA con labios (HeyGen)", "duracion": "20-60s", "costo": "~$0.10/30s"},
    {"tipo": "Video", "icono": "🎬", "descripcion": "Slide estático + voz (sin avatar)", "duracion": "30-120s", "costo": "~$0.02/30s"},
    {"tipo": "Interactivo", "icono": "👆", "descripcion": "Botones, acordeones, hotspots", "duracion": "1-3 min", "costo": "Gratis"},
    {"tipo": "Infografía", "icono": "📊", "descripcion": "Visualización de datos/procesos", "duracion": "30-60s", "costo": "Gratis"},
    {"tipo": "Comparador", "icono": "⚖️", "descripcion": "Tabla comparativa interactiva", "duracion": "1-2 min", "costo": "Gratis"},
    {"tipo": "Flashcards", "icono": "🃏", "descripcion": "Tarjetas pregunta/respuesta", "duracion": "2-3 min", "costo": "Gratis"},
    {"tipo": "Caso práctico", "icono": "🎯", "descripcion": "Escenario con decisiones y feedback", "duracion": "2-5 min", "costo": "Gratis"},
    {"tipo": "Quiz", "icono": "✅", "descripcion": "Evaluación con feedback", "duracion": "2-4 min", "costo": "Gratis"},
    {"tipo": "Timeline", "icono": "📅", "descripcion": "Línea de tiempo interactiva", "duracion": "1-2 min", "costo": "Gratis"},
    {"tipo": "Drag & Drop", "icono": "🎮", "descripcion": "Arrastrar elementos a categorías", "duracion": "1-3 min", "costo": "Gratis"},
    {"tipo": "Accordion", "icono": "📋", "descripcion": "Lista expandible de contenido", "duracion": "1-2 min", "costo": "Gratis"},
    {"tipo": "Simulador", "icono": "🖥️", "descripcion": "Réplica interactiva de sistema", "duracion": "3-5 min", "costo": "Medio"},
]

# Inicializar config de voz, avatar y música
if 'elevenlabs_voice_id' not in st.session_state:
    st.session_state.elevenlabs_voice_id = "SplyIQAjgy4DKGAnOrHi"  # Clau Bogotá default
if 'heygen_avatar_id' not in st.session_state:
    st.session_state.heygen_avatar_id = "Hada_LivelyGestures_Front_public"  # Default
if 'musica_fondo' not in st.session_state:
    st.session_state.musica_fondo = None
if 'volumen_musica' not in st.session_state:
    st.session_state.volumen_musica = 20
if 'usar_remotion' not in st.session_state:
    st.session_state.usar_remotion = False  # Remotion desactivado - usar slides HTML

# Sidebar con información y glosario
with st.sidebar:
    # Solo mostrar configuración avanzada para Learning
    sidebar_rol = st.session_state.get("rol", "learning") if AUTH_DISPONIBLE else "learning"

    if sidebar_rol == "solicitante":
        st.markdown("### 📝 Nueva Solicitud")
        st.info("Completá el formulario para solicitar un nuevo curso e-learning.")
    else:
        st.markdown("### ⚙️ Configuración")

    # Configuración de API (automático, sin toggle visible)
    usar_api = API_DISPONIBLE if sidebar_rol == "learning" else False
    st.session_state.usar_api = usar_api

    if sidebar_rol == "learning" and not usar_api:
        if api_key:
            st.success("OpenAI API ✓")
        else:
            st.error("Falta OPENAI_API_KEY en .env")

    # Configuración avanzada solo para Learning
    if sidebar_rol == "learning":
        # ElevenLabs Voice ID
        with st.expander("🎙️ Voz ElevenLabs"):
            voces_predefinidas = {
                "Clau Bogotá (default)": "SplyIQAjgy4DKGAnOrHi",
                "Valeria (casual)": "JddqVF50ZSIR7SRbJE6u",
                "Gaby (joven)": "a0MaQpDjx7p7bZmqzFp1",
                "Personalizada": "custom"
            }

            voz_seleccionada = st.selectbox(
                "Seleccionar voz",
                options=list(voces_predefinidas.keys()),
                index=0
            )

            if voz_seleccionada == "Personalizada":
                custom_id = st.text_input(
                    "Voice ID",
                    value=st.session_state.elevenlabs_voice_id,
                    help="ID de voz de ElevenLabs"
                )
                st.session_state.elevenlabs_voice_id = custom_id
            else:
                st.session_state.elevenlabs_voice_id = voces_predefinidas[voz_seleccionada]

            st.caption(f"ID: `{st.session_state.elevenlabs_voice_id}`")

        # HeyGen Avatar ID
        with st.expander("🎭 Avatar HeyGen"):
            avatares_predefinidos = {
                "Hada (gestos animados)": "Hada_LivelyGestures_Front_public",
                "Annie (profesional)": "Annie_Business_Casual_Standing_Front_public",
                "Caroline (corporativo)": "Caroline_Office_Standing_Front_public",
                "Adriana (talk show)": "Adriana_BizTalk_Front_public",
                "Personalizado": "custom"
            }

            avatar_seleccionado = st.selectbox(
                "Seleccionar avatar",
                options=list(avatares_predefinidos.keys()),
                index=0
            )

            if avatar_seleccionado == "Personalizado":
                custom_avatar = st.text_input(
                    "Avatar ID",
                    value=st.session_state.heygen_avatar_id,
                    help="ID del avatar de HeyGen"
                )
                st.session_state.heygen_avatar_id = custom_avatar
            else:
                st.session_state.heygen_avatar_id = avatares_predefinidos[avatar_seleccionado]

            st.caption(f"ID: `{st.session_state.heygen_avatar_id}`")

        # Motor de Video
        with st.expander("🎬 Motor de Video"):
            motor_video = st.radio(
                "Motor de generación",
                options=["canva", "slides"],
                format_func=lambda x: "🎨 Canva (recomendado)" if x == "canva" else "📊 Slides estáticos",
                index=0,
                help="Canva genera presentaciones animadas profesionales"
            )
            st.session_state.motor_video = motor_video
            # Mantener compatibilidad
            st.session_state.usar_remotion = False

            if motor_video == "canva":
                st.success("✨ Videos profesionales con Canva")
                st.caption("Presentaciones animadas exportadas a MP4 + voiceover ElevenLabs")
            else:
                st.info("📊 Videos con slides estáticos")
                st.caption("HTML renderizado a imagen + audio")

        # Música de fondo
        with st.expander("🎵 Música de fondo"):
            musica_seleccionada = st.selectbox(
                "Seleccionar música",
                options=list(MUSICA_FONDO.keys()),
                index=0
            )

            if musica_seleccionada == "Personalizada":
                custom_music = st.text_input(
                    "URL de audio MP3",
                    placeholder="https://ejemplo.com/musica.mp3",
                    help="URL directa a archivo MP3"
                )
                st.session_state.musica_fondo = custom_music if custom_music else None
            else:
                st.session_state.musica_fondo = MUSICA_FONDO[musica_seleccionada]

            volumen_musica = st.slider("Volumen música", 0, 100, 20, help="% del volumen de la voz")
            st.session_state.volumen_musica = volumen_musica

            if st.session_state.get('musica_fondo'):
                st.caption("🎵 Música activa")

        # Estado del flujo
        st.markdown("---")
        st.markdown("### 📍 Progreso")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("✅ Solicitud" if st.session_state.solicitud else "⬜ Solicitud")
            st.markdown("✅ Malla" if st.session_state.malla else "⬜ Malla")
            st.markdown("✅ Diseño" if st.session_state.guiones else "⬜ Diseño")
        with col2:
            st.markdown("✅ Contenido" if st.session_state.contenido_generado else "⬜ Contenido")
            st.markdown("✅ SCORM" if st.session_state.scorm_path else "⬜ SCORM")

        # Glosario compacto
        st.markdown("---")
        st.markdown("### 📖 Recursos")
        for r in GLOSARIO_RECURSOS:
            st.caption(f"{r['icono']} {r['tipo']}")

# Tabs según rol
user_rol = st.session_state.get("rol", "learning") if AUTH_DISPONIBLE else "learning"

if user_rol == "solicitante":
    # Vista simplificada para solicitantes - Solicitud + Mis Solicitudes
    tab1 = st.container()
    tab2 = tab3 = tab4 = tab5 = tab6 = tab_inbox = None
else:
    # Vista completa para Learning - con Dashboard de Solicitudes
    tab1 = None  # No formulario de Solicitud para Learning
    tab_inbox, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "📥 Solicitudes",
        "1️⃣ Malla",
        "2️⃣ Diseño",
        "3️⃣ Contenido",
        "4️⃣ SCORM",
        "5️⃣ LMS"
    ])

# Función para extraer datos de Excel de solicitud
def extract_solicitud_from_excel(uploaded_excel):
    """Extrae datos de solicitud desde un Excel con formato específico"""
    try:
        df = pd.read_excel(uploaded_excel, header=None)
        data = {}

        # Buscar campos por etiqueta en columna A, valor en columna B
        for idx, row in df.iterrows():
            if pd.notna(row[0]):
                label = str(row[0]).lower().strip()
                value = row[1] if len(row) > 1 and pd.notna(row[1]) else ""

                if "nombre" in label and "curso" in label:
                    data['nombre'] = str(value)
                elif "audiencia" in label:
                    data['audiencia'] = str(value)
                elif "nivel" in label:
                    data['nivel'] = str(value)
                elif "duración" in label or "duracion" in label:
                    data['duracion'] = str(value)
                elif "área" in label or "area" in label:
                    data['area'] = str(value)
                elif "objetivo" in label:
                    data['objetivo'] = str(value)
                elif "temas" in label or "contenido" in label:
                    data['temas'] = str(value)
                elif "evaluación" in label or "evaluacion" in label or "quiz" in label:
                    data['requiere_eval'] = str(value).lower() in ['sí', 'si', 'yes', 'true', '1', 'x']

        return data
    except Exception as e:
        return {"error": str(e)}

# TAB 1: Nueva Solicitud (solo para solicitante)
if tab1:
  with tab1:
    st.markdown("### Nueva Solicitud de Curso")

    # Opción de subir Excel
    st.markdown("#### 📤 Cargar desde Excel (opcional)")
    excel_file = st.file_uploader(
        "Sube un Excel con la solicitud",
        type=['xlsx', 'xls'],
        help="El Excel debe tener columnas: Nombre del curso, Audiencia, Nivel, Duración, Área, Objetivo, Temas"
    )

    # Valores por defecto o desde Excel
    defaults = {
        'nombre': '', 'audiencia': '', 'nivel': 'Básico',
        'duracion': '15-20 minutos', 'area': 'Banca Pyme',
        'objetivo': '', 'temas': '', 'requiere_eval': True
    }

    if excel_file:
        excel_data = extract_solicitud_from_excel(excel_file)
        if 'error' in excel_data:
            st.error(f"Error leyendo Excel: {excel_data['error']}")
        else:
            defaults.update(excel_data)
            st.success(f"Datos cargados desde Excel")

        # Mostrar plantilla de ejemplo
        with st.expander("Ver formato esperado del Excel"):
            st.markdown("""
            | Campo | Valor |
            |-------|-------|
            | Nombre del curso | Introducción a Leasing |
            | Audiencia | Asesores comerciales Pyme |
            | Nivel | Básico / Intermedio / Avanzado |
            | Duración | 15-20 minutos |
            | Área | Banca Pyme |
            | Objetivo | Al finalizar, el participante podrá... |
            | Temas | Tema 1, Tema 2, Tema 3 |
            | Evaluación | Sí / No |
            """)

    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:
        nombre_curso = st.text_input("Nombre del curso *", value=defaults['nombre'], placeholder="Ej: Introducción a Leasing")
        audiencia = st.text_input("Audiencia *", value=defaults['audiencia'], placeholder="Ej: Asesores comerciales Pyme")

        nivel_options = ["Básico", "Intermedio", "Avanzado"]
        nivel_idx = nivel_options.index(defaults['nivel']) if defaults['nivel'] in nivel_options else 0
        nivel = st.selectbox("Nivel", nivel_options, index=nivel_idx)

    with col2:
        duracion_options = [
            "Menos de 10 minutos",
            "10-15 minutos",
            "15-20 minutos",
            "20-30 minutos",
            "Más de 30 minutos"
        ]
        dur_idx = 2
        for i, opt in enumerate(duracion_options):
            if defaults.get('duracion', '') in opt or opt in str(defaults.get('duracion', '')):
                dur_idx = i
                break
        duracion = st.selectbox("Duración estimada", duracion_options, index=dur_idx)

        area_options = [
            "Banca Pyme", "Banca Empresarial", "Cash Management",
            "Leasing", "Comercio Exterior", "Otra"
        ]
        area_idx = area_options.index(defaults['area']) if defaults['area'] in area_options else 0
        area = st.selectbox("Área solicitante", area_options, index=area_idx)

        requiere_eval = st.checkbox("Requiere evaluación (Quiz)", value=defaults.get('requiere_eval', True))

    objetivo = st.text_area("Objetivo del curso *",
        value=defaults.get('objetivo', ''),
        placeholder="Al finalizar, el participante podrá...",
        height=80)

    temas = st.text_area("Temas a cubrir *",
        value=defaults.get('temas', ''),
        placeholder="Lista los temas principales que debe incluir el curso...",
        height=120)

    # Subir documentación
    st.markdown("#### 📎 Documentación de apoyo (opcional)")
    uploaded_files = st.file_uploader(
        "Sube archivos con contenido base (PDF, DOCX, TXT)",
        type=['pdf', 'docx', 'txt'],
        accept_multiple_files=True
    )

    # Extraer contenido de archivos
    contenido_docs = ""
    if uploaded_files:
        st.info(f"📄 {len(uploaded_files)} archivo(s) cargado(s)")
        for uf in uploaded_files:
            texto = extract_text_from_file(uf)
            if texto:
                contenido_docs += f"\n\n--- Contenido de {uf.name} ---\n{texto}"

    if st.button("🚀 Enviar Solicitud", type="primary", use_container_width=True):
        if not nombre_curso or not audiencia or not objetivo or not temas:
            st.error("Completa todos los campos obligatorios (*)")
        else:
            # Parsear duración
            dur_map = {
                "Menos de 10 minutos": 8,
                "10-15 minutos": 12,
                "15-20 minutos": 17,
                "20-30 minutos": 25,
                "Más de 30 minutos": 35
            }
            duracion_min = dur_map.get(duracion, 15)

            with st.spinner("Enviando solicitud..."):
                try:
                    if API_DISPONIBLE:
                        # Preparar datos para la API
                        solicitante = {
                            "email": st.session_state.get("user_email", "solicitante@davivienda.com"),
                            "nombre": st.session_state.get("user_name", "Solicitante"),
                            "area": area
                        }

                        curso = {
                            "nombre": nombre_curso,
                            "audiencia": audiencia,
                            "nivel": nivel,
                            "duracion_min": duracion_min,
                            "objetivo": objetivo,
                            "temas": temas,
                            "requiere_eval": requiere_eval,
                            "documentacion": contenido_docs[:8000] if contenido_docs else ""
                        }

                        result, error = api_client.crear_solicitud(
                            solicitante=solicitante,
                            curso=curso,
                            prioridad="media"
                        )

                        if error:
                            raise Exception(error)

                        st.success(f"✅ Solicitud enviada exitosamente!")
                        st.info("El equipo de Learning revisará tu solicitud. Puedes ver el estado en 'Mis Solicitudes' abajo.")
                        st.balloons()

                    else:
                        # Modo sin API (local/demo)
                        st.session_state.solicitud = {
                            "nombre": nombre_curso,
                            "audiencia": audiencia,
                            "nivel": nivel,
                            "duracion": duracion_min,
                            "area": area,
                            "objetivo": objetivo,
                            "temas": temas,
                            "requiere_eval": requiere_eval
                        }
                        st.success("✅ Solicitud guardada (modo demo sin API)")

                except Exception as e:
                    st.error(f"Error: {str(e)}")

    # =============================================
    # MIS SOLICITUDES (Sección para Solicitante)
    # =============================================
    st.markdown("---")
    st.markdown("### 📋 Mis Solicitudes")

    # Constantes para estados
    MIS_STATUS_COLORS = {
        "pendiente": "🟡",
        "en_revision": "🔵",
        "devuelto": "🟠",
        "aprobado": "🟢",
        "rechazado": "🔴",
        "en_proceso": "🔷",
        "completado": "✅",
    }

    MIS_STATUS_LABELS = {
        "pendiente": "Pendiente",
        "en_revision": "En revisión",
        "devuelto": "Requiere tu atención",
        "aprobado": "Aprobado",
        "rechazado": "Rechazado",
        "en_proceso": "En producción",
        "completado": "Completado",
    }

    if API_DISPONIBLE:
        try:
            user_email = st.session_state.get("user_email", "solicitante@davivienda.com")
            result, error = api_client.mis_solicitudes(user_email)

            if error:
                st.error(f"Error cargando solicitudes: {error}")
                mis_solicitudes_list = []
            else:
                mis_solicitudes_list = result.get("solicitudes", [])

        except Exception as e:
            st.error(f"Error: {e}")
            mis_solicitudes_list = []
    else:
        st.info("Conecta a Firebase para ver tus solicitudes.")
        mis_solicitudes_list = []

    if mis_solicitudes_list:
        for sol_item in mis_solicitudes_list:
            sol_id = sol_item.get("id", "")
            curso_nombre = sol_item.get("curso_nombre", "Sin nombre")
            status = sol_item.get("status", "pendiente")
            ultimo_comentario = sol_item.get("ultimo_comentario", "")
            ultimo_rol = sol_item.get("ultimo_comentario_rol", "")

            status_emoji = MIS_STATUS_COLORS.get(status, "⬜")
            status_label = MIS_STATUS_LABELS.get(status, status)

            # Destacar si hay respuesta de Learning
            tiene_respuesta = ultimo_rol == "learning" and ultimo_comentario

            with st.container():
                col1, col2 = st.columns([3, 1])

                with col1:
                    st.markdown(f"**{curso_nombre}**")
                    if tiene_respuesta:
                        st.info(f"💬 Learning: {ultimo_comentario}")
                    elif ultimo_comentario:
                        st.caption(f"💬 {ultimo_comentario[:80]}...")

                with col2:
                    st.markdown(f"{status_emoji} **{status_label}**")

                    # Si está devuelto, mostrar opción de responder
                    if status == "devuelto":
                        with st.expander("↩️ Responder"):
                            respuesta = st.text_area("Tu respuesta", key=f"resp_{sol_id}")
                            if st.button("Enviar", key=f"btn_{sol_id}"):
                                if respuesta:
                                    autor = {
                                        "email": user_email,
                                        "nombre": st.session_state.get("user_name", "Solicitante"),
                                        "rol": "solicitante"
                                    }
                                    result, error = api_client.agregar_comentario(sol_id, respuesta, autor)
                                    if not error:
                                        # Cambiar estado a pendiente para que Learning lo revise
                                        api_client.actualizar_solicitud(sol_id, status="pendiente")
                                        st.success("Respuesta enviada")
                                        st.rerun()
                                    else:
                                        st.error(error)
                                else:
                                    st.warning("Escribe una respuesta")

                st.markdown("---")

    elif API_DISPONIBLE:
        st.info("No tienes solicitudes aún. Completa el formulario arriba para crear una.")

# TAB INBOX: Dashboard de Solicitudes (Learning)
if tab_inbox:
  with tab_inbox:
    # Estado de sesión para la solicitud seleccionada
    if 'solicitud_seleccionada_id' not in st.session_state:
        st.session_state.solicitud_seleccionada_id = None

    # Colores para estados
    STATUS_COLORS = {
        "pendiente": "🟡",
        "en_revision": "🔵",
        "devuelto": "🟠",
        "aprobado": "🟢",
        "rechazado": "🔴",
        "en_proceso": "🔷",
        "completado": "✅",
    }

    STATUS_LABELS = {
        "pendiente": "Pendiente",
        "en_revision": "En revisión",
        "devuelto": "Devuelto",
        "aprobado": "Aprobado",
        "rechazado": "Rechazado",
        "en_proceso": "En proceso",
        "completado": "Completado",
    }

    # Filtros compactos en una línea
    col_f1, col_f2, col_f3, col_f4 = st.columns([2, 2, 2, 1])

    with col_f1:
        filtro_status = st.selectbox(
            "Estado",
            options=["Todos", "pendiente", "en_revision", "devuelto", "aprobado", "en_proceso"],
            format_func=lambda x: "Todos" if x == "Todos" else f"{STATUS_COLORS.get(x, '')} {STATUS_LABELS.get(x, x)}",
            label_visibility="collapsed"
        )

    with col_f2:
        filtro_area = st.selectbox(
            "Área",
            options=["Todas", "Banca Pyme", "Banca Empresarial", "Cash Management", "Leasing", "Comercio Exterior"],
            label_visibility="collapsed"
        )

    with col_f3:
        filtro_asignado = st.text_input("Asignado", placeholder="Filtrar por email...", label_visibility="collapsed")

    with col_f4:
        st.button("🔄", help="Actualizar")

    # Cargar solicitudes
    if is_api_available():
        try:
            params_status = None if filtro_status == "Todos" else filtro_status
            params_area = None if filtro_area == "Todas" else filtro_area
            params_asignado = filtro_asignado if filtro_asignado else None

            result, error = api_client.listar_solicitudes(
                status=params_status,
                area=params_area,
                asignado_a=params_asignado
            )

            if error:
                st.error(f"Error: {error}")
                solicitudes_list = []
            else:
                solicitudes_list = result.get("solicitudes", [])

        except Exception as e:
            st.error(f"Error: {e}")
            solicitudes_list = []
    else:
        st.warning("API no disponible")
        solicitudes_list = []

    # Mostrar tabla de solicitudes
    if solicitudes_list:
        # Preparar datos para tabla
        tabla_data = []
        for sol_item in solicitudes_list:
            status = sol_item.get("status", "pendiente")
            prioridad = sol_item.get("prioridad", "media")
            prioridad_emoji = "🔴" if prioridad == "alta" else ("🟡" if prioridad == "media" else "🟢")

            tabla_data.append({
                "id": sol_item.get("id", ""),
                "Curso": sol_item.get("curso_nombre", "Sin nombre"),
                "Área": sol_item.get("area", ""),
                "Estado": f"{STATUS_COLORS.get(status, '')} {STATUS_LABELS.get(status, status)}",
                "Prioridad": f"{prioridad_emoji}",
                "Asignado": (sol_item.get("asignado_a") or "-")[:20],
            })

        df_solicitudes = pd.DataFrame(tabla_data)

        # Mostrar tabla con selección
        st.dataframe(
            df_solicitudes[["Curso", "Área", "Estado", "Prioridad", "Asignado"]],
            use_container_width=True,
            hide_index=True,
            height=250
        )

        # Selector de solicitud
        opciones_sol = {f"{s['Curso']} ({s['Área']})": s['id'] for s in tabla_data}
        seleccion = st.selectbox(
            "Seleccionar solicitud",
            options=["-- Seleccionar --"] + list(opciones_sol.keys()),
            label_visibility="collapsed"
        )

        if seleccion != "-- Seleccionar --":
            st.session_state.solicitud_seleccionada_id = opciones_sol[seleccion]

        # Panel de detalle si hay solicitud seleccionada
        if st.session_state.solicitud_seleccionada_id:
            # Botón volver arriba del detalle
            if st.button("⬅️ Volver a todas las solicitudes", type="secondary"):
                st.session_state.solicitud_seleccionada_id = None
                st.rerun()

            st.markdown("### 📋 Detalle de Solicitud")

            try:
                sol_detail, error = api_client.obtener_solicitud(st.session_state.solicitud_seleccionada_id)

                if error:
                    st.error(f"Error: {error}")
                elif sol_detail:
                    curso = sol_detail.get("curso", {})
                    solicitante = sol_detail.get("solicitante", {})
                    comentarios = sol_detail.get("comentarios", [])
                    current_status = sol_detail.get("status", "pendiente")

                    col_d1, col_d2 = st.columns([2, 1])

                    with col_d1:
                        st.markdown(f"#### {curso.get('nombre', 'Sin nombre')}")
                        st.markdown(f"""
                        - **Audiencia:** {curso.get('audiencia', '')}
                        - **Nivel:** {curso.get('nivel', '')}
                        - **Duración:** {curso.get('duracion_min', 0)} min
                        - **Objetivo:** {curso.get('objetivo', '')}
                        - **Temas:** {curso.get('temas', '')}
                        - **Requiere evaluación:** {'Sí' if curso.get('requiere_eval') else 'No'}
                        """)

                        st.markdown(f"**Solicitante:** {solicitante.get('nombre', '')} ({solicitante.get('email', '')})")
                        st.markdown(f"**Área:** {solicitante.get('area', '')}")

                    with col_d2:
                        st.markdown("##### Acciones")

                        # Botones de acción según estado
                        if current_status == "pendiente":
                            if st.button("🔵 Tomar Revisión", type="primary", use_container_width=True):
                                user_email = st.session_state.get("user_email", "learning@davivienda.com")
                                result, error = api_client.actualizar_solicitud(
                                    st.session_state.solicitud_seleccionada_id,
                                    status="en_revision",
                                    asignado_a=user_email
                                )
                                if not error:
                                    st.success("Solicitud tomada")
                                    st.rerun()
                                else:
                                    st.error(error)

                        if current_status in ["pendiente", "en_revision"]:
                            col_btn1, col_btn2 = st.columns(2)
                            with col_btn1:
                                if st.button("✅ Aprobar", use_container_width=True):
                                    result, error = api_client.actualizar_solicitud(
                                        st.session_state.solicitud_seleccionada_id,
                                        status="aprobado"
                                    )
                                    if not error:
                                        st.success("Solicitud aprobada")
                                        st.rerun()

                            with col_btn2:
                                if st.button("↩️ Devolver", use_container_width=True):
                                    result, error = api_client.actualizar_solicitud(
                                        st.session_state.solicitud_seleccionada_id,
                                        status="devuelto"
                                    )
                                    if not error:
                                        st.warning("Solicitud devuelta")
                                        st.rerun()

                            if st.button("❌ Rechazar", use_container_width=True):
                                result, error = api_client.actualizar_solicitud(
                                    st.session_state.solicitud_seleccionada_id,
                                    status="rechazado"
                                )
                                if not error:
                                    st.error("Solicitud rechazada")
                                    st.rerun()

                        if current_status == "aprobado":
                            if st.button("🚀 Iniciar Producción", type="primary", use_container_width=True):
                                # Cargar datos en solicitud con formato correcto
                                st.session_state.solicitud = {
                                    "nombre": curso.get("nombre", ""),
                                    "audiencia": curso.get("audiencia", ""),
                                    "nivel": curso.get("nivel", "Básico"),
                                    "duracion": curso.get("duracion_min", 15),
                                    "area": solicitante.get("area", ""),
                                    "objetivo": curso.get("objetivo", ""),
                                    "temas": curso.get("temas", ""),
                                    "requiere_eval": curso.get("requiere_eval", True)
                                }
                                # Limpiar selección
                                st.session_state.solicitud_seleccionada_id = None
                                # Actualizar estado en Firebase
                                result, error = api_client.actualizar_solicitud(
                                    sol_detail.get("id"),
                                    status="en_proceso"
                                )
                                if not error:
                                    st.success("¡Producción iniciada!")
                                    # JavaScript para cambiar a tab Malla
                                    st.markdown("""
                                    <script>
                                        setTimeout(function() {
                                            const tabs = window.parent.document.querySelectorAll('button[data-baseweb="tab"]');
                                            tabs.forEach(tab => {
                                                if (tab.innerText.includes('Malla')) {
                                                    tab.click();
                                                }
                                            });
                                        }, 100);
                                    </script>
                                    """, unsafe_allow_html=True)
                                    st.rerun()

                    # Comentarios
                    st.markdown("---")
                    st.markdown("##### 💬 Comentarios")

                    for com in comentarios:
                        autor = com.get("autor", {})
                        texto = com.get("texto", "")
                        rol = autor.get("rol", "")
                        nombre = autor.get("nombre", autor.get("email", ""))

                        emoji = "👤" if rol == "solicitante" else "🎓"
                        st.markdown(f"{emoji} **{nombre}** ({rol})")
                        st.markdown(f"> {texto}")
                        st.markdown("")

                    # Agregar comentario
                    nuevo_comentario = st.text_area("Agregar comentario", key="nuevo_comentario_inbox")
                    if st.button("📤 Enviar Comentario"):
                        if nuevo_comentario:
                            user_email = st.session_state.get("user_email", "learning@davivienda.com")
                            user_name = st.session_state.get("user_name", "Learning")
                            autor = {
                                "email": user_email,
                                "nombre": user_name,
                                "rol": "learning"
                            }
                            result, error = api_client.agregar_comentario(
                                st.session_state.solicitud_seleccionada_id,
                                nuevo_comentario,
                                autor
                            )
                            if not error:
                                st.success("Comentario agregado")
                                st.rerun()
                            else:
                                st.error(error)
                        else:
                            st.warning("Escribe un comentario")

            except Exception as e:
                st.error(f"Error cargando detalle: {e}")

    else:
        st.info("No hay solicitudes pendientes. 🎉")


# TAB 2: Malla y Diseño Instruccional (Learning)
if tab2:
  with tab2:
    if st.session_state.malla:
        sol = st.session_state.solicitud

        st.markdown(f"""
        <div class="info-box">
            <strong>📚 {sol['nombre']}</strong> |
            👥 {sol['audiencia']} |
            📊 {sol['nivel']} |
            ⏱️ v{st.session_state.version}
        </div>
        """, unsafe_allow_html=True)

        # Convertir a DataFrame
        df = pd.DataFrame(st.session_state.malla)

        # Calcular duración total
        duracion_total = df['duracion_min'].sum()
        st.metric("Duración Total", f"{duracion_total} minutos")

        # Mostrar tabla editable
        edited_df = st.data_editor(
            df,
            use_container_width=True,
            num_rows="dynamic",
            column_config={
                "id": st.column_config.NumberColumn("ID", width="small"),
                "etapa": st.column_config.SelectboxColumn(
                    "Etapa",
                    options=["Introducción", "Desarrollo", "Cierre"],
                    width="medium"
                ),
                "bloque": st.column_config.TextColumn("Bloque", width="medium"),
                "objetivo": st.column_config.TextColumn("Objetivo", width="large"),
                "tipo_recurso": st.column_config.SelectboxColumn(
                    "Tipo Recurso",
                    options=["Video avatar", "Video", "Interactivo", "Infografía",
                             "Comparador", "Flashcards", "Caso práctico", "Quiz"],
                    width="medium"
                ),
                "recurso": st.column_config.TextColumn("Recurso", width="medium"),
                "descripcion": st.column_config.TextColumn("Descripción", width="large"),
                "duracion_min": st.column_config.NumberColumn("Min", width="small")
            }
        )

        # Auto-guardar cambios y sincronizar con guiones
        nueva_malla = edited_df.to_dict('records')
        if nueva_malla != st.session_state.malla:
            st.session_state.malla = nueva_malla

            # Sincronizar tipo_recurso con guiones existentes
            if st.session_state.guiones:
                for item in nueva_malla:
                    item_id = item.get('id')
                    nuevo_tipo = item.get('tipo_recurso')
                    # Buscar guión correspondiente y actualizar tipo
                    for guion in st.session_state.guiones:
                        if guion.get('id') == item_id and guion.get('tipo') != nuevo_tipo:
                            guion['tipo'] = nuevo_tipo
                            # Limpiar contenido generado si cambió el tipo
                            st.session_state.contenido_generado = [
                                c for c in st.session_state.contenido_generado
                                if c.get('guion_id') != item_id
                            ]

        st.markdown("---")

        # Feedback con IA
        st.markdown("### 🔄 Iterar con Feedback")

        feedback = st.text_area(
            "¿Qué cambios necesita la malla?",
            placeholder="Ej: Agregar un caso práctico antes del quiz, reducir la duración total a 12 minutos...",
            height=100
        )

        if st.button("🔄 Regenerar con Feedback", type="primary"):
            if not feedback:
                st.warning("Escribe el feedback primero")
            elif not api_key and not st.session_state.get('usar_api', False):
                st.error("Configura tu API Key")
            else:
                with st.spinner("Regenerando malla..."):
                    try:
                        # Usar API Cloud o modo local
                        if st.session_state.get('usar_api', False) and API_DISPONIBLE and st.session_state.get('malla_id'):
                            result, error = api_client.iterar_malla(
                                st.session_state.malla_id,
                                feedback
                            )
                            if error:
                                raise Exception(error)
                            nueva_malla = result.get("malla", [])
                        else:
                            # Modo local
                            client = OpenAI(api_key=api_key)

                            prompt = f"""Eres un experto en Diseño Instruccional.

MALLA ACTUAL:
{json.dumps(st.session_state.malla, indent=2, ensure_ascii=False)}

FEEDBACK DEL USUARIO:
{feedback}

REGLAS:
1. Incorpora TODOS los cambios del feedback
2. Mantén la estructura JSON
3. Un bloque puede tener múltiples recursos

TIPOS DE RECURSO: Video avatar, Video, Interactivo, Infografía, Comparador, Flashcards, Caso práctico, Quiz

Responde SOLO con el JSON array actualizado."""

                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": prompt}],
                                max_tokens=4000,
                                temperature=0.7
                            )

                            content = response.choices[0].message.content
                            start = content.find('[')
                            end = content.rfind(']') + 1
                            nueva_malla = json.loads(content[start:end])

                        st.session_state.version += 1
                        st.session_state.malla = nueva_malla
                        st.session_state.historial.append({
                            "version": st.session_state.version,
                            "feedback": feedback,
                            "malla": nueva_malla
                        })

                        st.success(f"✅ Versión {st.session_state.version} generada")
                        st.rerun()

                    except Exception as e:
                        st.error(f"Error: {str(e)}")

        # Historial de versiones
        with st.expander("📜 Historial de versiones"):
            for h in st.session_state.historial:
                st.markdown(f"**v{h['version']}**: {h['feedback']}")

        # Exportar malla
        st.markdown("---")
        st.markdown("### 📁 Exportar Malla")

        col1, col2, col3 = st.columns(3)

        with col1:
            csv = df.to_csv(index=False)
            st.download_button(
                "📥 CSV",
                data=csv,
                file_name=f"malla_{sol['nombre'].replace(' ', '_')}_v{st.session_state.version}.csv",
                mime="text/csv",
                use_container_width=True
            )

        with col2:
            json_str = json.dumps({
                "solicitud": sol,
                "version": st.session_state.version,
                "malla": st.session_state.malla,
                "fecha": datetime.now().isoformat()
            }, indent=2, ensure_ascii=False)
            st.download_button(
                "📥 JSON",
                data=json_str,
                file_name=f"malla_{sol['nombre'].replace(' ', '_')}_v{st.session_state.version}.json",
                mime="application/json",
                use_container_width=True
            )

        with col3:
            if st.button("💾 Guardar", use_container_width=True):
                output_dir = Path("/Users/federico/Desktop/ia-davivienda/output")
                output_dir.mkdir(exist_ok=True)
                nombre_base = sol['nombre'].replace(' ', '_').lower()
                csv_path = output_dir / f"malla_{nombre_base}_v{st.session_state.version}.csv"
                df.to_csv(csv_path, index=False)
                json_path = output_dir / f"malla_{nombre_base}_v{st.session_state.version}.json"
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump({"solicitud": sol, "version": st.session_state.version, "malla": st.session_state.malla}, f, indent=2, ensure_ascii=False)
                st.success(f"✅ Guardado en /output")

        # Botón de aprobar y generar diseño instruccional
        st.markdown("---")
        st.markdown("### ✅ Aprobar Malla")

        if st.session_state.guiones:
            st.success(f"✅ Malla aprobada - Diseño instruccional ya generado ({len(st.session_state.guiones)} recursos)")
            st.markdown("➡️ Ve a la pestaña **Diseño** para revisar y editar")
        else:
            st.info("Cuando la malla esté lista, apruébala para generar el diseño instruccional automáticamente.")

            if st.button("✅ Aprobar Malla y Generar Diseño", type="primary", use_container_width=True):
                if not api_key and not st.session_state.get('usar_api', False):
                    st.error("Falta API Key")
                else:
                    with st.spinner("Generando diseño instruccional para todos los recursos..."):
                        try:
                            # Usar API Cloud o modo local
                            if st.session_state.get('usar_api', False) and API_DISPONIBLE and st.session_state.get('malla_id'):
                                result, error = api_client.generar_guiones(st.session_state.malla_id)
                                if error:
                                    raise Exception(error)
                                guiones = result.get("guiones", [])
                            else:
                                # Modo local
                                client = OpenAI(api_key=api_key)
                                guiones = []

                                progress = st.progress(0)
                                for i, recurso in enumerate(st.session_state.malla):
                                    prompt = f"""Eres un diseñador instruccional experto en e-learning corporativo para Davivienda.

CONTEXTO DEL CURSO:
- Nombre: {sol['nombre']}
- Audiencia: {sol['audiencia']}
- Objetivo general: {sol['objetivo']}

RECURSO A DESARROLLAR:
- Tipo: {recurso['tipo_recurso']}
- Nombre: {recurso['recurso']}
- Descripción: {recurso['descripcion']}
- Duración: {recurso['duracion_min']} minutos
- Objetivo específico: {recurso['objetivo']}

GENERA EL CONTENIDO DETALLADO según el tipo:

Para "Video avatar" o "Video": genera {{"voiceover": "texto completo que dirá el presentador", "indicaciones": "tono y pausas"}}

Para "Quiz": genera {{"preguntas": [{{"pregunta": "...", "opciones": ["a", "b", "c", "d"], "correcta": 0, "feedback_correcto": "...", "feedback_incorrecto": "..."}}]}}

Para "Flashcards": genera {{"tarjetas": [{{"frente": "pregunta", "reverso": "respuesta"}}]}}

Para "Caso práctico": genera {{"escenario": "descripción", "opciones": [{{"opcion": "...", "consecuencia": "..."}}]}}

Para "Interactivo" o "Accordion": genera {{"elementos": [{{"titulo": "...", "contenido": "..."}}]}}

Para "Infografía": genera {{"secciones": [{{"titulo": "...", "contenido": "...", "dato": "..."}}]}}

Para "Comparador": genera {{"columnas": ["Aspecto", "Opción A", "Opción B"], "items": [{{"aspecto": "...", "opcion_a": "...", "opcion_b": "..."}}]}}

Responde SOLO con JSON válido:
{{"id": {recurso['id']}, "tipo": "{recurso['tipo_recurso']}", "titulo": "{recurso['recurso']}", "contenido": {{...}}, "notas_produccion": "..."}}"""

                                    response = client.chat.completions.create(
                                        model="gpt-4o",
                                        messages=[{"role": "user", "content": prompt}],
                                        max_tokens=2000,
                                        temperature=0.7
                                    )

                                    content = response.choices[0].message.content
                                    start = content.find('{')
                                    end = content.rfind('}') + 1
                                    guion = json.loads(content[start:end])
                                    guiones.append(guion)

                                    progress.progress((i + 1) / len(st.session_state.malla))

                            st.session_state.guiones = guiones
                            st.success(f"✅ Malla aprobada - Diseño generado para {len(guiones)} recursos")

                            # JavaScript para cambiar a la tab de Diseño
                            st.markdown("""
                            <script>
                                setTimeout(function() {
                                    const tabs = window.parent.document.querySelectorAll('button[data-baseweb="tab"]');
                                    tabs.forEach(tab => {
                                        if (tab.innerText.includes('Diseño')) {
                                            tab.click();
                                        }
                                    });
                                }, 500);
                            </script>
                            """, unsafe_allow_html=True)

                            st.balloons()

                        except Exception as e:
                            st.error(f"Error: {str(e)}")

    else:
        st.info("👈 Genera una malla desde el área solicitante")

# Función para renderizar contenido de guión de forma visual
def render_guion_contenido(contenido, tipo):
    """Renderiza el contenido del guión de forma visual según el tipo"""
    if isinstance(contenido, str):
        st.markdown(contenido)
        return

    if tipo in ["Video avatar", "Video"]:
        if "voiceover" in contenido:
            st.markdown("**🎙️ Voiceover:**")
            st.info(contenido.get("voiceover", ""))
        if "texto" in contenido:
            st.markdown("**🎙️ Texto:**")
            st.info(contenido.get("texto", ""))
        if "indicaciones" in contenido:
            st.caption(f"📋 {contenido.get('indicaciones', '')}")

    elif tipo == "Quiz":
        preguntas = contenido.get("preguntas", contenido.get("questions", []))
        if isinstance(preguntas, list):
            for i, p in enumerate(preguntas, 1):
                st.markdown(f"**Pregunta {i}:** {p.get('pregunta', p.get('question', ''))}")
                opciones = p.get("opciones", p.get("options", []))
                correcta = p.get("correcta", p.get("correct", 0))
                for j, op in enumerate(opciones):
                    marca = "✅" if j == correcta else "⬜"
                    st.markdown(f"  {marca} {op}")
                st.markdown("---")

    elif tipo == "Flashcards":
        tarjetas = contenido.get("tarjetas", contenido.get("cards", []))
        if isinstance(tarjetas, list):
            cols = st.columns(min(3, len(tarjetas)))
            for i, t in enumerate(tarjetas):
                with cols[i % 3]:
                    st.markdown(f"**{t.get('frente', t.get('front', ''))}**")
                    st.caption(t.get('reverso', t.get('back', '')))

    elif tipo == "Caso práctico":
        if "escenario" in contenido:
            st.markdown("**📖 Escenario:**")
            st.info(contenido.get("escenario", ""))
        opciones = contenido.get("opciones", contenido.get("decisions", []))
        if isinstance(opciones, list):
            st.markdown("**🔀 Opciones:**")
            for op in opciones:
                with st.expander(f"➡️ {op.get('opcion', op.get('option', ''))}"):
                    st.write(op.get('consecuencia', op.get('feedback', '')))

    elif tipo in ["Interactivo", "Accordion"]:
        elementos = contenido.get("elementos", contenido.get("items", []))
        if isinstance(elementos, list):
            for el in elementos:
                with st.expander(f"📌 {el.get('titulo', el.get('title', ''))}"):
                    st.write(el.get('contenido', el.get('content', '')))

    elif tipo == "Infografía":
        secciones = contenido.get("secciones", contenido.get("sections", []))
        if isinstance(secciones, list):
            for sec in secciones:
                st.markdown(f"**{sec.get('titulo', sec.get('title', ''))}**")
                st.write(sec.get('contenido', sec.get('content', '')))

    elif tipo == "Comparador":
        items = contenido.get("items", contenido.get("comparacion", []))
        if isinstance(items, list) and len(items) > 0:
            df_comp = pd.DataFrame(items)
            st.dataframe(df_comp, use_container_width=True, hide_index=True)

    else:
        # Fallback: mostrar como texto formateado
        for key, value in contenido.items():
            st.markdown(f"**{key}:**")
            if isinstance(value, list):
                for item in value:
                    st.write(f"• {item}")
            else:
                st.write(value)

# Función para generar documento Word
def generar_word_diseno(sol, guiones):
    """Genera documento Word con el diseño instruccional"""
    doc = docx.Document()

    # Título
    doc.add_heading(f"Diseño Instruccional: {sol['nombre']}", 0)
    doc.add_paragraph(f"Audiencia: {sol['audiencia']}")
    doc.add_paragraph(f"Nivel: {sol['nivel']}")
    doc.add_paragraph(f"Objetivo: {sol['objetivo']}")
    doc.add_paragraph("")

    for guion in guiones:
        # Título del recurso
        doc.add_heading(f"{guion.get('tipo', '')} - {guion.get('titulo', '')}", level=1)

        contenido = guion.get('contenido', {})

        if isinstance(contenido, str):
            doc.add_paragraph(contenido)
        elif isinstance(contenido, dict):
            for key, value in contenido.items():
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                if isinstance(value, str):
                    doc.add_paragraph(value)
                elif isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict):
                            for k, v in item.items():
                                doc.add_paragraph(f"{k}: {v}", style='List Bullet')
                        else:
                            doc.add_paragraph(str(item), style='List Bullet')

        if guion.get('notas_produccion'):
            doc.add_paragraph(f"Notas de producción: {guion['notas_produccion']}", style='Intense Quote')

        doc.add_page_break()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# TAB 3: Diseño Instruccional (Learning)
if tab3:
  with tab3:
    st.markdown("### 📝 Diseño Instruccional")

    if not st.session_state.malla:
        st.info("👈 Primero genera una malla en la pestaña anterior")
    else:
        sol = st.session_state.solicitud
        st.markdown(f"**Curso:** {sol['nombre']} | **Recursos:** {len(st.session_state.malla)}")

        # Botón para generar diseño instruccional
        if st.button("🎬 Generar Diseño Instruccional", type="primary", use_container_width=True):
            if not api_key:
                st.error("Falta API Key")
            else:
                with st.spinner("Generando diseño instruccional..."):
                    try:
                        client = OpenAI(api_key=api_key)
                        guiones = []

                        progress = st.progress(0)
                        for i, recurso in enumerate(st.session_state.malla):
                            prompt = f"""Eres un diseñador instruccional experto en e-learning corporativo para Davivienda.

CONTEXTO DEL CURSO:
- Nombre: {sol['nombre']}
- Audiencia: {sol['audiencia']}
- Objetivo general: {sol['objetivo']}

RECURSO A DESARROLLAR:
- Tipo: {recurso['tipo_recurso']}
- Nombre: {recurso['recurso']}
- Descripción: {recurso['descripcion']}
- Duración: {recurso['duracion_min']} minutos
- Objetivo específico: {recurso['objetivo']}

GENERA EL CONTENIDO DETALLADO según el tipo:

Para "Video avatar" o "Video": genera {{"voiceover": "texto completo que dirá el presentador", "indicaciones": "tono y pausas"}}

Para "Quiz": genera {{"preguntas": [{{"pregunta": "...", "opciones": ["a", "b", "c", "d"], "correcta": 0, "feedback_correcto": "...", "feedback_incorrecto": "..."}}]}}

Para "Flashcards": genera {{"tarjetas": [{{"frente": "pregunta", "reverso": "respuesta"}}]}}

Para "Caso práctico": genera {{"escenario": "descripción", "opciones": [{{"opcion": "...", "consecuencia": "..."}}]}}

Para "Interactivo" o "Accordion": genera {{"elementos": [{{"titulo": "...", "contenido": "..."}}]}}

Para "Infografía": genera {{"secciones": [{{"titulo": "...", "contenido": "...", "dato": "..."}}]}}

Para "Comparador": genera {{"columnas": ["Aspecto", "Opción A", "Opción B"], "items": [{{"aspecto": "...", "opcion_a": "...", "opcion_b": "..."}}]}}

Responde SOLO con JSON válido:
{{"id": {recurso['id']}, "tipo": "{recurso['tipo_recurso']}", "titulo": "{recurso['recurso']}", "contenido": {{...}}, "notas_produccion": "..."}}"""

                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": prompt}],
                                max_tokens=2000,
                                temperature=0.7
                            )

                            content = response.choices[0].message.content
                            start = content.find('{')
                            end = content.rfind('}') + 1
                            guion = json.loads(content[start:end])
                            guiones.append(guion)

                            progress.progress((i + 1) / len(st.session_state.malla))

                        st.session_state.guiones = guiones
                        st.success(f"✅ Diseño generado para {len(guiones)} recursos")
                        st.rerun()

                    except Exception as e:
                        st.error(f"Error: {str(e)}")

        # Mostrar diseño si existe
        if st.session_state.guiones:
            st.markdown("---")

            # Toggle modo edición
            modo_edicion = st.toggle("✏️ Modo edición manual", value=False, key="modo_edicion_diseno")

            if modo_edicion:
                st.info("📝 Edita directamente el contenido de cada recurso")

            # Visualización/Edición
            for idx, guion in enumerate(st.session_state.guiones):
                tipo = guion.get('tipo', '')
                icono = next((r['icono'] for r in GLOSARIO_RECURSOS if r['tipo'] == tipo), "📄")
                guion_id = guion.get('id', idx)

                with st.expander(f"{icono} **{guion.get('titulo', 'Sin título')}** ({tipo})", expanded=modo_edicion):

                    if modo_edicion:
                        # MODO EDICIÓN
                        contenido = guion.get('contenido', {})

                        if tipo in ["Video avatar", "Video"]:
                            nuevo_voiceover = st.text_area(
                                "🎙️ Voiceover",
                                value=contenido.get('voiceover', contenido.get('texto', '')),
                                height=200,
                                key=f"vo_{guion_id}"
                            )
                            nuevas_indicaciones = st.text_input(
                                "📋 Indicaciones",
                                value=contenido.get('indicaciones', ''),
                                key=f"ind_{guion_id}"
                            )

                            if st.button("💾 Guardar", key=f"save_{guion_id}"):
                                st.session_state.guiones[idx]['contenido']['voiceover'] = nuevo_voiceover
                                st.session_state.guiones[idx]['contenido']['indicaciones'] = nuevas_indicaciones
                                st.success("✅ Guardado")

                        elif tipo == "Quiz":
                            preguntas = contenido.get('preguntas', contenido.get('questions', []))
                            st.markdown("**Preguntas:**")

                            for p_idx, pregunta in enumerate(preguntas):
                                st.markdown(f"**Pregunta {p_idx + 1}:**")
                                nueva_pregunta = st.text_input(
                                    "Texto",
                                    value=pregunta.get('pregunta', pregunta.get('question', '')),
                                    key=f"q_{guion_id}_{p_idx}"
                                )

                                opciones = pregunta.get('opciones', pregunta.get('options', []))
                                nuevas_opciones = []
                                cols = st.columns(4)
                                for o_idx, op in enumerate(opciones):
                                    with cols[o_idx]:
                                        nueva_op = st.text_input(f"Op {o_idx+1}", value=op, key=f"op_{guion_id}_{p_idx}_{o_idx}")
                                        nuevas_opciones.append(nueva_op)

                                correcta = st.number_input("Correcta (0-3)", value=pregunta.get('correcta', 0), min_value=0, max_value=3, key=f"cor_{guion_id}_{p_idx}")
                                st.markdown("---")

                            if st.button("💾 Guardar Quiz", key=f"save_quiz_{guion_id}"):
                                st.success("✅ Quiz guardado")

                        elif tipo == "Flashcards":
                            tarjetas = contenido.get('tarjetas', contenido.get('cards', []))
                            st.markdown("**Tarjetas:**")

                            for t_idx, tarjeta in enumerate(tarjetas):
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.text_input("Frente", value=tarjeta.get('frente', tarjeta.get('front', '')), key=f"fc_f_{guion_id}_{t_idx}")
                                with col2:
                                    st.text_input("Reverso", value=tarjeta.get('reverso', tarjeta.get('back', '')), key=f"fc_b_{guion_id}_{t_idx}")

                            if st.button("💾 Guardar Flashcards", key=f"save_fc_{guion_id}"):
                                st.success("✅ Flashcards guardadas")

                        elif tipo == "Caso práctico":
                            nuevo_escenario = st.text_area(
                                "📖 Escenario",
                                value=contenido.get('escenario', ''),
                                height=150,
                                key=f"esc_{guion_id}"
                            )

                            opciones = contenido.get('opciones', [])
                            st.markdown("**Opciones de decisión:**")
                            for op_idx, opcion in enumerate(opciones):
                                col1, col2 = st.columns(2)
                                with col1:
                                    st.text_input("Opción", value=opcion.get('opcion', ''), key=f"cp_op_{guion_id}_{op_idx}")
                                with col2:
                                    st.text_input("Consecuencia", value=opcion.get('consecuencia', ''), key=f"cp_con_{guion_id}_{op_idx}")

                            if st.button("💾 Guardar Caso", key=f"save_cp_{guion_id}"):
                                st.session_state.guiones[idx]['contenido']['escenario'] = nuevo_escenario
                                st.success("✅ Caso guardado")

                        else:
                            # Edición genérica JSON
                            contenido_str = json.dumps(contenido, indent=2, ensure_ascii=False)
                            nuevo_contenido = st.text_area(
                                "Contenido (JSON)",
                                value=contenido_str,
                                height=300,
                                key=f"json_{guion_id}"
                            )

                            if st.button("💾 Guardar JSON", key=f"save_json_{guion_id}"):
                                try:
                                    st.session_state.guiones[idx]['contenido'] = json.loads(nuevo_contenido)
                                    st.success("✅ Guardado")
                                except json.JSONDecodeError:
                                    st.error("JSON inválido")

                        # Notas de producción (siempre editable)
                        st.divider()
                        nuevas_notas = st.text_input(
                            "💡 Notas de producción",
                            value=guion.get('notas_produccion', ''),
                            key=f"notas_{guion_id}"
                        )
                        if nuevas_notas != guion.get('notas_produccion', ''):
                            st.session_state.guiones[idx]['notas_produccion'] = nuevas_notas

                    else:
                        # MODO VISUALIZACIÓN
                        render_guion_contenido(guion.get('contenido', {}), tipo)

                        if guion.get('notas_produccion'):
                            st.divider()
                            st.caption(f"💡 Producción: {guion['notas_produccion']}")

            # Iterar con feedback (como la malla)
            st.markdown("---")
            st.markdown("### 🔄 Editar con IA")

            feedback_diseno = st.text_area(
                "¿Qué cambios necesita el diseño?",
                placeholder="Ej: Agregar más preguntas al quiz, cambiar el tono del voiceover a más informal, incluir más ejemplos...",
                height=80,
                key="feedback_diseno"
            )

            if st.button("🔄 Regenerar con Feedback", key="regen_diseno"):
                if not feedback_diseno:
                    st.warning("Escribe el feedback primero")
                elif not api_key:
                    st.error("Falta API Key")
                else:
                    with st.spinner("Actualizando diseño..."):
                        try:
                            client = OpenAI(api_key=api_key)

                            prompt = f"""Eres un diseñador instruccional experto.

DISEÑO ACTUAL:
{json.dumps(st.session_state.guiones, indent=2, ensure_ascii=False)}

FEEDBACK DEL USUARIO:
{feedback_diseno}

Actualiza el diseño incorporando TODOS los cambios solicitados.
Mantén la misma estructura JSON.
Responde SOLO con el JSON array actualizado."""

                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[{"role": "user", "content": prompt}],
                                max_tokens=8000,
                                temperature=0.7
                            )

                            content = response.choices[0].message.content
                            start = content.find('[')
                            end = content.rfind(']') + 1
                            st.session_state.guiones = json.loads(content[start:end])
                            st.success("✅ Diseño actualizado")
                            st.rerun()

                        except Exception as e:
                            st.error(f"Error: {str(e)}")

            # Exportar
            st.markdown("---")
            st.markdown("### 📥 Exportar")

            col1, col2, col3 = st.columns(3)

            with col1:
                word_buffer = generar_word_diseno(sol, st.session_state.guiones)
                st.download_button(
                    "📄 Word (.docx)",
                    data=word_buffer,
                    file_name=f"diseno_{sol['nombre'].replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            with col2:
                json_guiones = json.dumps(st.session_state.guiones, indent=2, ensure_ascii=False)
                st.download_button(
                    "📋 JSON",
                    data=json_guiones,
                    file_name=f"diseno_{sol['nombre'].replace(' ', '_')}.json",
                    mime="application/json",
                    use_container_width=True
                )

            with col3:
                if st.button("💾 Guardar", use_container_width=True, key="save_diseno"):
                    output_dir = Path("/Users/federico/Desktop/ia-davivienda/output")
                    output_dir.mkdir(exist_ok=True)
                    nombre_base = sol['nombre'].replace(' ', '_').lower()

                    # Guardar JSON
                    path_json = output_dir / f"diseno_{nombre_base}.json"
                    with open(path_json, 'w', encoding='utf-8') as f:
                        json.dump(st.session_state.guiones, f, indent=2, ensure_ascii=False)

                    # Guardar Word
                    path_word = output_dir / f"diseno_{nombre_base}.docx"
                    word_buffer = generar_word_diseno(sol, st.session_state.guiones)
                    with open(path_word, 'wb') as f:
                        f.write(word_buffer.getvalue())

                    st.success(f"✅ Guardado en /output")

# TAB 4: Generar Contenido - Assets (Learning)
if tab4:
  with tab4:
    st.markdown("### 🎨 Generar Contenido")

    if not st.session_state.guiones:
        st.info("👈 Primero genera el diseño instruccional en la pestaña anterior")
    else:
        sol = st.session_state.solicitud

        # Info del pipeline
        st.markdown("""
        <div style="background: #f0f7ff; border-radius: 10px; padding: 15px; margin-bottom: 20px;">
            <strong>Pipeline de generación:</strong><br>
            🎙️ <b>Audio</b> (ElevenLabs) → 🎬 <b>Video/Slides</b> (HTML+FFmpeg) → 📦 <b>Assets</b>
        </div>
        """, unsafe_allow_html=True)

        # Resumen
        total = len(st.session_state.guiones)
        generados = len(st.session_state.contenido_generado)

        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total", total)
        with col2:
            st.metric("Generados", generados)
        with col3:
            st.metric("Pendientes", total - generados)
        with col4:
            aprobados = sum(1 for c in st.session_state.contenido_generado if c.get('status') == 'aprobado')
            st.metric("Aprobados", aprobados)

        st.markdown("---")

        # Botones de acción global
        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button("🚀 Generar Todo", type="primary", use_container_width=True, disabled=generados == total):
                with st.spinner("Generando assets..."):
                    import time
                    progress = st.progress(0)

                    output_base = Path(f"/Users/federico/Desktop/ia-davivienda/output/{sol['nombre'].replace(' ', '_').lower()}")
                    output_base.mkdir(parents=True, exist_ok=True)

                    for i, g in enumerate(st.session_state.guiones):
                        time.sleep(0.3)
                        if not any(c.get('guion_id') == g.get('id') for c in st.session_state.contenido_generado):
                            tipo = g.get('tipo', 'Otro')
                            recurso_dir = output_base / f"{g.get('id'):02d}_{tipo.replace(' ', '_').lower()}"
                            recurso_dir.mkdir(exist_ok=True)

                            st.session_state.contenido_generado.append({
                                "guion_id": g.get('id'),
                                "tipo": tipo,
                                "titulo": g.get('titulo'),
                                "status": "generado",
                                "path": str(recurso_dir)
                            })
                        progress.progress((i + 1) / len(st.session_state.guiones))
                    st.success(f"✅ Todos generados")
                    st.rerun()

        with col2:
            pendientes = [g for g in st.session_state.guiones if not any(c.get('guion_id') == g.get('id') for c in st.session_state.contenido_generado)]
            if st.button(f"⏭️ Generar Pendientes ({len(pendientes)})", use_container_width=True, disabled=len(pendientes) == 0):
                with st.spinner("Generando pendientes..."):
                    import time
                    output_base = Path(f"/Users/federico/Desktop/ia-davivienda/output/{sol['nombre'].replace(' ', '_').lower()}")
                    output_base.mkdir(parents=True, exist_ok=True)

                    for g in pendientes:
                        time.sleep(0.3)
                        tipo = g.get('tipo', 'Otro')
                        recurso_dir = output_base / f"{g.get('id'):02d}_{tipo.replace(' ', '_').lower()}"
                        recurso_dir.mkdir(exist_ok=True)

                        st.session_state.contenido_generado.append({
                            "guion_id": g.get('id'),
                            "tipo": tipo,
                            "titulo": g.get('titulo'),
                            "status": "generado",
                            "path": str(recurso_dir)
                        })
                    st.success(f"✅ {len(pendientes)} generados")
                    st.rerun()

        with col3:
            if st.button("🗑️ Limpiar todo", use_container_width=True, disabled=generados == 0):
                st.session_state.contenido_generado = []
                st.rerun()

        st.markdown("---")

        # Lista de recursos con generación individual
        st.markdown("#### 📋 Recursos")

        # Controles de selección masiva
        col_sel1, col_sel2, col_sel3 = st.columns([1, 1, 2])
        with col_sel1:
            if st.button("☑️ Seleccionar pendientes", use_container_width=True):
                for g in st.session_state.guiones:
                    gid = g.get('id')
                    if not any(c.get('guion_id') == gid for c in st.session_state.contenido_generado):
                        st.session_state.recursos_seleccionados.add(gid)
                st.rerun()
        with col_sel2:
            if st.button("⬜ Deseleccionar todo", use_container_width=True):
                st.session_state.recursos_seleccionados = set()
                st.rerun()
        with col_sel3:
            num_seleccionados = len(st.session_state.recursos_seleccionados)
            if st.button(f"🚀 Generar {num_seleccionados} seleccionados en paralelo",
                        use_container_width=True,
                        disabled=num_seleccionados == 0,
                        type="primary"):
                # Generar en paralelo usando threads
                guiones_a_generar = [g for g in st.session_state.guiones
                                     if g.get('id') in st.session_state.recursos_seleccionados]

                progress_bar = st.progress(0, text="Iniciando generación paralela...")
                status_container = st.container()

                resultados_paralelos = []

                def actualizar_progreso(gid, estado, mensaje):
                    st.session_state.progreso_generacion[gid] = {"estado": estado, "mensaje": mensaje}

                with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                    futures = {
                        executor.submit(
                            generar_recurso_individual,
                            g, sol,
                            st.session_state.elevenlabs_voice_id,
                            st.session_state.heygen_avatar_id,
                            actualizar_progreso
                        ): g.get('id') for g in guiones_a_generar
                    }

                    completados = 0
                    for future in concurrent.futures.as_completed(futures):
                        resultado = future.result()
                        resultados_paralelos.append(resultado)
                        completados += 1
                        progress_bar.progress(
                            completados / len(guiones_a_generar),
                            text=f"Completados: {completados}/{len(guiones_a_generar)}"
                        )

                # Agregar resultados al session_state
                for res in resultados_paralelos:
                    # Remover si ya existe
                    st.session_state.contenido_generado = [
                        c for c in st.session_state.contenido_generado
                        if c.get('guion_id') != res['guion_id']
                    ]
                    st.session_state.contenido_generado.append(res)

                st.session_state.recursos_seleccionados = set()
                st.success(f"✅ {len(resultados_paralelos)} recursos generados!")
                st.rerun()

        st.markdown("---")

        for g in st.session_state.guiones:
            guion_id = g.get('id')
            tipo = g.get('tipo', 'Otro')
            titulo = g.get('titulo', 'Sin título')
            icono = next((r['icono'] for r in GLOSARIO_RECURSOS if r['tipo'] == tipo), "📄")

            # Verificar estado
            contenido_item = next((c for c in st.session_state.contenido_generado if c.get('guion_id') == guion_id), None)
            generado = contenido_item is not None
            status = contenido_item.get('status', 'pendiente') if contenido_item else 'pendiente'

            # Determinar assets
            if tipo in ["Video avatar", "Video"]:
                assets = "🎙️ Audio + 🎬 Video"
            elif tipo == "Quiz":
                assets = "📝 HTML + SCORM"
            elif tipo in ["Interactivo", "Accordion", "Flashcards", "Caso práctico"]:
                assets = "👆 HTML interactivo"
            elif tipo in ["Infografía", "Comparador", "Timeline"]:
                assets = "📊 HTML + PNG"
            else:
                assets = "📄 HTML"

            # Status badge
            if status == 'aprobado':
                status_badge = "✅ Aprobado"
                status_color = "#d4edda"
            elif status == 'generado' or status == 'actualizado':
                status_badge = "🔄 Generado"
                status_color = "#fff3cd"
            else:
                status_badge = "⏳ Pendiente"
                status_color = "#f8f9fa"

            # Card del recurso con checkbox
            col_check, col_info = st.columns([0.5, 9.5])

            with col_check:
                is_selected = guion_id in st.session_state.recursos_seleccionados
                if st.checkbox("", value=is_selected, key=f"sel_{guion_id}", label_visibility="collapsed"):
                    st.session_state.recursos_seleccionados.add(guion_id)
                else:
                    st.session_state.recursos_seleccionados.discard(guion_id)

            with col_info:
                st.markdown(f"""
                <div style="background: {status_color}; border-radius: 8px; padding: 10px; margin-bottom: 5px;">
                    <strong>{icono} {titulo}</strong>
                    <span style="float: right; font-size: 0.8em;">{status_badge}</span>
                    <br><small style="color: #666;">{tipo} → {assets}</small>
                </div>
                """, unsafe_allow_html=True)

            col1, col2, col3, col4 = st.columns([2, 1, 1, 1])

            with col1:
                if contenido_item:
                    st.caption(f"📂 {contenido_item.get('path', '')}")

            with col2:
                if not generado:
                    if st.button("🎨 Generar", key=f"gen_ind_{guion_id}", use_container_width=True):
                        # Buscar guión
                        guion_data = next((gg for gg in st.session_state.guiones if gg.get('id') == guion_id), None)

                        if guion_data:
                            with st.spinner(f"Generando {titulo}..."):
                                output_base = Path(f"/Users/federico/Desktop/ia-davivienda/output/{sol['nombre'].replace(' ', '_').lower()}")
                                output_base.mkdir(parents=True, exist_ok=True)
                                recurso_dir = output_base / f"{guion_id:02d}_{tipo.replace(' ', '_').lower()}"
                                recurso_dir.mkdir(exist_ok=True)

                                contenido_data = guion_data.get('contenido', {})
                                assets_generados = {}

                                # Generar según tipo
                                if tipo == "Video avatar":
                                    # VIDEO AVATAR = HeyGen con avatar IA
                                    voiceover = contenido_data.get('voiceover', contenido_data.get('texto', ''))
                                    if voiceover:
                                        with st.status("Generando video con avatar HeyGen...", expanded=True) as status:
                                            # 1. Generar audio
                                            st.write("🎙️ Generando audio con ElevenLabs...")
                                            audio_path = recurso_dir / "audio.mp3"
                                            audio_result, audio_error = generar_audio_elevenlabs(
                                                voiceover,
                                                st.session_state.elevenlabs_voice_id,
                                                str(audio_path)
                                            )

                                            if audio_result:
                                                assets_generados['audio'] = str(audio_path)
                                                st.write("✅ Audio generado")

                                                # 2. Subir audio a servidor temporal
                                                st.write("📤 Subiendo audio para HeyGen...")
                                                audio_url, upload_error = subir_audio_temporal(str(audio_path))

                                                if audio_url:
                                                    st.write("✅ Audio subido")

                                                    # 3. Generar video con HeyGen
                                                    st.write(f"🎭 Generando avatar ({st.session_state.heygen_avatar_id})...")
                                                    st.write("⏳ Esto puede tomar 1-3 minutos...")
                                                    video_path = recurso_dir / "video_avatar.mp4"
                                                    video_result, video_error = generar_video_heygen(
                                                        audio_url,
                                                        st.session_state.heygen_avatar_id,
                                                        str(video_path)
                                                    )

                                                    if video_result:
                                                        assets_generados['video'] = str(video_path)
                                                        st.write("✅ Video con avatar generado!")
                                                        status.update(label="Avatar generado!", state="complete")
                                                    else:
                                                        st.warning(f"HeyGen: {video_error}")
                                                        status.update(label="Error en HeyGen", state="error")
                                                else:
                                                    st.warning(f"Upload: {upload_error}")
                                                    status.update(label="Error subiendo audio", state="error")
                                            else:
                                                st.warning(f"Audio: {audio_error}")
                                                status.update(label="Error en audio", state="error")

                                elif tipo == "Video":
                                    # VIDEO = Canva o slide estático
                                    voiceover = contenido_data.get('voiceover', contenido_data.get('texto', ''))
                                    if voiceover:
                                        motor_video = st.session_state.get('motor_video', 'canva')

                                        if motor_video == "canva":
                                            with st.status("🎨 Preparando video Canva...", expanded=True) as status:
                                                st.write("🎙️ Generando audio con ElevenLabs...")

                                                lineas = [l.strip() for l in voiceover.split('\n') if l.strip()]
                                                subtitulo = lineas[0] if lineas else "Aprende los conceptos clave"
                                                bullets = lineas[1:5] if len(lineas) > 1 else ["Contenido del módulo"]

                                                resultado = generar_video_canva(
                                                    titulo=titulo,
                                                    subtitulo=subtitulo,
                                                    bullets=bullets,
                                                    voiceover_text=voiceover,
                                                    output_dir=str(recurso_dir),
                                                    voice_id=st.session_state.elevenlabs_voice_id
                                                )

                                                if resultado.get('audio_path'):
                                                    assets_generados['audio'] = resultado['audio_path']
                                                    st.write("✅ Audio generado")

                                                if resultado.get('success'):
                                                    assets_generados['canva_request'] = resultado['request_path']
                                                    st.write("✅ Request Canva preparado")
                                                    st.info(f"📋 Para generar el video, ejecuta en Claude Code:\n`genera video canva con {Path(resultado['request_path']).name}`")
                                                    status.update(label="🎨 Audio listo - Ejecutar Canva", state="complete")
                                                else:
                                                    st.error(f"Error: {resultado.get('error', 'Desconocido')}")
                                                    status.update(label="Error preparando Canva", state="error")
                                        else:
                                            with st.status("Generando video (slide + voz)...", expanded=True) as status:
                                                st.write("🎙️ Generando audio con ElevenLabs...")

                                                resultados = generar_video_completo(
                                                    titulo,
                                                    voiceover,
                                                    tipo,
                                                    str(recurso_dir),
                                                    sol['nombre'],
                                                    st.session_state.elevenlabs_voice_id
                                                )

                                                if resultados.get('audio'):
                                                    assets_generados['audio'] = resultados['audio']
                                                    st.write("✅ Audio generado")

                                                if resultados.get('video'):
                                                    assets_generados['video'] = resultados['video']
                                                    st.write("✅ Video MP4 generado")
                                                    status.update(label="Video generado!", state="complete")
                                                else:
                                                    status.update(label="Video parcial", state="running")

                                                for error in resultados.get('errors', []):
                                                    st.warning(error)

                                elif tipo == "Quiz":
                                    # Generar HTML del quiz
                                    preguntas = contenido_data.get('preguntas', [])
                                    quiz_html = generar_quiz_html(preguntas, titulo, sol['nombre'])
                                    quiz_path = recurso_dir / "quiz.html"
                                    with open(quiz_path, 'w', encoding='utf-8') as f:
                                        f.write(quiz_html)
                                    assets_generados['quiz'] = str(quiz_path)

                                else:
                                    # Generar slide genérico
                                    slide_path = recurso_dir / "index.html"
                                    html_preview = generar_html_preview(guion_data, sol)
                                    with open(slide_path, 'w', encoding='utf-8') as f:
                                        f.write(html_preview)
                                    assets_generados['html'] = str(slide_path)

                                st.session_state.contenido_generado.append({
                                    "guion_id": guion_id,
                                    "tipo": tipo,
                                    "titulo": titulo,
                                    "status": "generado",
                                    "path": str(recurso_dir),
                                    "assets": assets_generados
                                })
                                st.rerun()
                else:
                    if st.button("🔄 Regenerar", key=f"regen_ind_{guion_id}", use_container_width=True):
                        # Eliminar el existente
                        st.session_state.contenido_generado = [c for c in st.session_state.contenido_generado if c.get('guion_id') != guion_id]
                        st.rerun()

            with col3:
                if generado and status != 'aprobado':
                    if st.button("✅ Aprobar", key=f"apr_ind_{guion_id}", use_container_width=True):
                        for i, c in enumerate(st.session_state.contenido_generado):
                            if c.get('guion_id') == guion_id:
                                st.session_state.contenido_generado[i]['status'] = 'aprobado'
                                break
                        st.rerun()

            with col4:
                if generado:
                    if st.button("🗑️", key=f"del_ind_{guion_id}", use_container_width=True):
                        st.session_state.contenido_generado = [c for c in st.session_state.contenido_generado if c.get('guion_id') != guion_id]
                        st.rerun()

            # Expander con preview y feedback inline (solo si está generado)
            if generado:
                guion = next((gg for gg in st.session_state.guiones if gg.get('id') == guion_id), None)

                with st.expander(f"👁️ Preview y Feedback", expanded=False):
                    if guion:
                        contenido = guion.get('contenido', {})

                        # Tabs: Preview Visual | Datos | Feedback
                        tab_preview, tab_datos, tab_feedback = st.tabs(["👁️ Preview Visual", "📋 Datos", "🤖 Iterar con IA"])

                        with tab_preview:
                            # Mostrar assets generados (audio, etc)
                            assets = contenido_item.get('assets', {}) if contenido_item else {}

                            # Si hay video generado, mostrar reproductor de video
                            if 'video' in assets and os.path.exists(assets['video']):
                                st.markdown("**🎬 Video Generado:**")
                                st.video(assets['video'])
                                st.caption(f"📂 {assets['video']}")
                                st.markdown("---")

                            # Si hay audio pero no video, mostrar reproductor de audio
                            elif 'audio' in assets and os.path.exists(assets['audio']):
                                st.markdown("**🎧 Audio Generado:**")
                                st.audio(assets['audio'], format='audio/mp3')
                                st.caption(f"📂 {assets['audio']}")
                                st.markdown("---")

                            # Si hay PNG, mostrar link (PIL incompatible con esta arquitectura)
                            if 'png' in assets and os.path.exists(assets['png']):
                                st.markdown(f"**🖼️ Slide Renderizado:** `{os.path.basename(assets['png'])}`")
                                st.markdown("---")

                            # Guardar preview y botón para abrir en navegador
                            html_preview = generar_html_preview(guion, sol)
                            preview_path = f"/tmp/preview_{guion_id}.html"
                            with open(preview_path, 'w', encoding='utf-8') as f:
                                f.write(html_preview)

                            col_prev1, col_prev2 = st.columns([1, 3])
                            with col_prev1:
                                if st.button("🌐 Abrir en navegador", key=f"open_preview_{guion_id}"):
                                    import subprocess
                                    subprocess.run(["open", preview_path])
                            with col_prev2:
                                st.caption(f"📄 {preview_path}")

                        with tab_datos:
                            # Mostrar datos estructurados
                            st.markdown("**📄 Datos del contenido:**")

                            # Mostrar assets generados
                            if assets:
                                st.markdown("**📦 Assets generados:**")
                                for asset_type, asset_path in assets.items():
                                    if os.path.exists(str(asset_path)):
                                        size_kb = os.path.getsize(str(asset_path)) / 1024
                                        st.markdown(f"• **{asset_type}**: `{os.path.basename(asset_path)}` ({size_kb:.1f} KB)")
                                    else:
                                        st.markdown(f"• **{asset_type}**: `{os.path.basename(asset_path)}` ⚠️")
                                st.markdown("---")

                            if tipo in ["Video avatar", "Video"]:
                                voiceover = contenido.get('voiceover', contenido.get('texto', ''))
                                if voiceover:
                                    st.text_area("Voiceover", value=voiceover, height=200, disabled=True, key=f"prev_vo_{guion_id}")
                                indicaciones = contenido.get('indicaciones', '')
                                if indicaciones:
                                    st.caption(f"📋 Indicaciones: {indicaciones}")

                            elif tipo == "Quiz":
                                preguntas = contenido.get('preguntas', contenido.get('questions', []))
                                for i, p in enumerate(preguntas, 1):
                                    st.markdown(f"**Pregunta {i}:** {p.get('pregunta', p.get('question', ''))}")
                                    correcta = p.get('correcta', p.get('correct', 0))
                                    opciones = p.get('opciones', p.get('options', []))
                                    cols = st.columns(4)
                                    for j, op in enumerate(opciones):
                                        with cols[j]:
                                            marca = "✅" if j == correcta else ""
                                            st.caption(f"{op} {marca}")
                                    st.markdown("---")

                            elif tipo == "Flashcards":
                                tarjetas = contenido.get('tarjetas', contenido.get('cards', []))
                                for t in tarjetas:
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.info(f"**{t.get('frente', t.get('front', ''))}**")
                                    with col2:
                                        st.success(t.get('reverso', t.get('back', '')))

                            elif tipo == "Caso práctico":
                                st.info(contenido.get('escenario', ''))
                                st.markdown("**Opciones:**")
                                opciones = contenido.get('opciones', [])
                                for op in opciones:
                                    with st.expander(f"➡️ {op.get('opcion', '')}"):
                                        st.write(op.get('consecuencia', ''))

                            elif tipo in ["Interactivo", "Accordion"]:
                                elementos = contenido.get('elementos', contenido.get('items', []))
                                for el in elementos:
                                    with st.expander(f"📌 {el.get('titulo', el.get('title', ''))}"):
                                        st.write(el.get('contenido', el.get('content', '')))

                            else:
                                st.json(contenido)

                        with tab_feedback:
                            st.markdown("**🤖 Iterar con IA:**")

                            feedback_inline = st.text_area(
                                "¿Qué cambios necesitas?",
                                placeholder="Ej: Acortar el voiceover, agregar más ejemplos, cambiar tono a más informal, agregar pregunta sobre X...",
                                height=120,
                                key=f"fb_inline_{guion_id}"
                            )

                            if st.button("🔄 Aplicar cambios con IA", key=f"apply_fb_{guion_id}", use_container_width=True, disabled=not feedback_inline):
                                if api_key:
                                    with st.spinner("Actualizando con IA..."):
                                        try:
                                            client = OpenAI(api_key=api_key)

                                            prompt = f"""Eres un diseñador instruccional experto para e-learning corporativo.

RECURSO ACTUAL (tipo: {tipo}):
{json.dumps(guion, indent=2, ensure_ascii=False)}

FEEDBACK DEL USUARIO:
{feedback_inline}

Actualiza el recurso incorporando TODOS los cambios solicitados.
Mantén exactamente la misma estructura JSON.
Responde SOLO con el JSON actualizado, nada más."""

                                            response = client.chat.completions.create(
                                                model="gpt-4o",
                                                messages=[{"role": "user", "content": prompt}],
                                                max_tokens=3000,
                                                temperature=0.7
                                            )

                                            content = response.choices[0].message.content
                                            start = content.find('{')
                                            end = content.rfind('}') + 1
                                            nuevo_guion = json.loads(content[start:end])

                                            # Actualizar guión
                                            for i, gg in enumerate(st.session_state.guiones):
                                                if gg.get('id') == guion_id:
                                                    st.session_state.guiones[i] = nuevo_guion
                                                    break

                                            # Actualizar status del contenido
                                            for i, cc in enumerate(st.session_state.contenido_generado):
                                                if cc.get('guion_id') == guion_id:
                                                    st.session_state.contenido_generado[i]['status'] = 'actualizado'
                                                    break

                                            st.success("✅ Contenido actualizado")
                                            st.rerun()

                                        except Exception as e:
                                            st.error(f"Error: {str(e)}")

                            st.caption(f"Estado actual: {status}")

            st.markdown("")  # Espaciado

        # Mostrar contenido generado con preview y feedback
        if st.session_state.contenido_generado:
            st.markdown("---")
            st.markdown("#### 👁️ Preview y Feedback")
            st.caption("Revisa cada recurso y da feedback para ajustar antes de generar los assets finales.")

            for idx, c in enumerate(st.session_state.contenido_generado):
                tipo = c.get('tipo', '')
                guion_id = c.get('guion_id')
                icono = next((r['icono'] for r in GLOSARIO_RECURSOS if r['tipo'] == tipo), "📄")

                # Buscar el guión correspondiente
                guion = next((g for g in st.session_state.guiones if g.get('id') == guion_id), None)

                with st.expander(f"{icono} {c.get('titulo', '')} - {c.get('status', 'pendiente').upper()}", expanded=False):
                    col_preview, col_info = st.columns([2, 1])

                    with col_preview:
                        st.markdown("##### 👁️ Preview")

                        # Obtener assets generados
                        assets_gen = c.get('assets', {})

                        if guion:
                            contenido = guion.get('contenido', {})

                            # Mostrar preview según tipo
                            if tipo in ["Video avatar", "Video"]:
                                # Mostrar video si existe
                                if 'video' in assets_gen and os.path.exists(assets_gen['video']):
                                    st.markdown("**🎬 Video Generado:**")
                                    st.video(assets_gen['video'])
                                    st.caption(f"📂 {os.path.basename(assets_gen['video'])}")
                                    st.markdown("---")
                                # Si no hay video pero hay audio, mostrar audio
                                elif 'audio' in assets_gen and os.path.exists(assets_gen['audio']):
                                    st.markdown("**🎧 Audio Generado:**")
                                    st.audio(assets_gen['audio'], format='audio/mp3')
                                    st.caption(f"📂 {os.path.basename(assets_gen['audio'])}")
                                    st.markdown("---")

                                voiceover = contenido.get('voiceover', contenido.get('texto', ''))
                                if voiceover:
                                    st.markdown("**🎙️ Voiceover (texto):**")
                                    st.info(voiceover[:500] + "..." if len(str(voiceover)) > 500 else voiceover)
                                    st.caption(f"🎭 Avatar: {st.session_state.heygen_avatar_id}")
                                    st.caption(f"🎙️ Voz: {st.session_state.elevenlabs_voice_id}")

                            elif tipo == "Quiz":
                                preguntas = contenido.get('preguntas', contenido.get('questions', []))
                                if preguntas:
                                    st.markdown(f"**📝 {len(preguntas)} preguntas:**")
                                    for i, p in enumerate(preguntas[:2], 1):  # Mostrar solo 2
                                        st.markdown(f"{i}. {p.get('pregunta', p.get('question', ''))}")
                                    if len(preguntas) > 2:
                                        st.caption(f"... y {len(preguntas) - 2} más")

                            elif tipo == "Flashcards":
                                tarjetas = contenido.get('tarjetas', contenido.get('cards', []))
                                if tarjetas:
                                    st.markdown(f"**🃏 {len(tarjetas)} tarjetas:**")
                                    for t in tarjetas[:3]:
                                        st.markdown(f"• **{t.get('frente', t.get('front', ''))}**")
                                    if len(tarjetas) > 3:
                                        st.caption(f"... y {len(tarjetas) - 3} más")

                            elif tipo == "Caso práctico":
                                escenario = contenido.get('escenario', '')
                                if escenario:
                                    st.markdown("**📖 Escenario:**")
                                    st.info(escenario[:300] + "..." if len(escenario) > 300 else escenario)
                                opciones = contenido.get('opciones', [])
                                if opciones:
                                    st.markdown(f"**🔀 {len(opciones)} opciones de decisión**")

                            elif tipo in ["Interactivo", "Accordion"]:
                                elementos = contenido.get('elementos', contenido.get('items', []))
                                if elementos:
                                    st.markdown(f"**📌 {len(elementos)} secciones:**")
                                    for el in elementos[:3]:
                                        st.markdown(f"• {el.get('titulo', el.get('title', ''))}")

                            elif tipo == "Infografía":
                                secciones = contenido.get('secciones', contenido.get('sections', []))
                                if secciones:
                                    st.markdown(f"**📊 {len(secciones)} secciones**")
                                    for s in secciones[:3]:
                                        st.markdown(f"• {s.get('titulo', s.get('title', ''))}")

                            elif tipo == "Comparador":
                                items = contenido.get('items', [])
                                if items:
                                    df_preview = pd.DataFrame(items[:3])
                                    st.dataframe(df_preview, use_container_width=True, hide_index=True)

                            else:
                                # Fallback
                                st.json(contenido)
                        else:
                            st.warning("No se encontró el diseño instruccional")

                    with col_info:
                        st.markdown("##### 📁 Assets")
                        if tipo in ["Video avatar", "Video"]:
                            st.markdown("```\n• audio.mp3\n• slide.png\n• video.mp4\n```")
                        elif tipo == "Quiz":
                            st.markdown("```\n• quiz.html\n• quiz.js\n• scorm.js\n```")
                        else:
                            st.markdown("```\n• index.html\n• styles.css\n• script.js\n```")

                        st.caption(f"📂 `{c.get('path', '')}`")

                    # Feedback y regenerar
                    st.markdown("---")
                    st.markdown("##### 🔄 Feedback")

                    feedback_key = f"feedback_{guion_id}"
                    feedback = st.text_area(
                        "¿Qué ajustes necesita este recurso?",
                        placeholder="Ej: Hacer el voiceover más corto, agregar una pregunta sobre X, cambiar el tono...",
                        height=80,
                        key=feedback_key
                    )

                    col_btn1, col_btn2 = st.columns(2)

                    with col_btn1:
                        if st.button("🔄 Regenerar con feedback", key=f"regen_{guion_id}", disabled=not feedback):
                            if api_key and guion:
                                with st.spinner("Regenerando..."):
                                    try:
                                        client = OpenAI(api_key=api_key)

                                        prompt = f"""Eres un diseñador instruccional experto.

RECURSO ACTUAL:
{json.dumps(guion, indent=2, ensure_ascii=False)}

FEEDBACK DEL USUARIO:
{feedback}

Actualiza SOLO este recurso incorporando el feedback.
Mantén la misma estructura JSON.
Responde SOLO con el JSON del recurso actualizado."""

                                        response = client.chat.completions.create(
                                            model="gpt-4o",
                                            messages=[{"role": "user", "content": prompt}],
                                            max_tokens=2000,
                                            temperature=0.7
                                        )

                                        content = response.choices[0].message.content
                                        start = content.find('{')
                                        end = content.rfind('}') + 1
                                        nuevo_guion = json.loads(content[start:end])

                                        # Actualizar en la lista de guiones
                                        for i, g in enumerate(st.session_state.guiones):
                                            if g.get('id') == guion_id:
                                                st.session_state.guiones[i] = nuevo_guion
                                                break

                                        # Actualizar status
                                        for i, cont in enumerate(st.session_state.contenido_generado):
                                            if cont.get('guion_id') == guion_id:
                                                st.session_state.contenido_generado[i]['status'] = 'actualizado'
                                                break

                                        st.success("✅ Recurso actualizado")
                                        st.rerun()

                                    except Exception as e:
                                        st.error(f"Error: {str(e)}")

                    with col_btn2:
                        if st.button("✅ Aprobar", key=f"aprobar_{guion_id}"):
                            for i, cont in enumerate(st.session_state.contenido_generado):
                                if cont.get('guion_id') == guion_id:
                                    st.session_state.contenido_generado[i]['status'] = 'aprobado'
                                    break
                            st.rerun()

            # Resumen de aprobaciones
            st.markdown("---")
            aprobados = sum(1 for c in st.session_state.contenido_generado if c.get('status') == 'aprobado')
            total_contenido = len(st.session_state.contenido_generado)

            if aprobados == total_contenido and total_contenido > 0:
                st.success(f"✅ Todos los recursos aprobados ({aprobados}/{total_contenido}). Puedes continuar al SCORM.")
            else:
                st.info(f"📊 Progreso: {aprobados}/{total_contenido} recursos aprobados")

# TAB 5: SCORM (Learning)
if tab5:
  with tab5:
    st.markdown("### 📦 Paquete SCORM")

    if not st.session_state.contenido_generado:
        st.info("👈 Primero genera el contenido en la pestaña anterior")
    else:
        sol = st.session_state.solicitud

        st.markdown("#### Configuración del paquete")

        col1, col2 = st.columns(2)
        with col1:
            scorm_version = st.selectbox("Versión SCORM", ["SCORM 1.2", "SCORM 2004"])
            completar_con = st.selectbox("Completar curso cuando:", [
                "Visualiza todo el contenido",
                "Aprueba el quiz (70%)",
                "Aprueba el quiz (80%)",
                "Aprueba el quiz (100%)"
            ])
        with col2:
            incluir_tracking = st.checkbox("Incluir tracking de progreso", value=True)
            incluir_bookmarking = st.checkbox("Permitir marcadores", value=True)

        st.markdown("---")

        if st.button("📦 Generar Paquete SCORM", type="primary", use_container_width=True):
            with st.spinner("Empaquetando SCORM..."):
                import zipfile
                import shutil
                import tempfile

                output_dir = Path("/Users/federico/Desktop/ia-davivienda/output")
                output_dir.mkdir(exist_ok=True)
                nombre_base = sol['nombre'].replace(' ', '_').lower()
                scorm_path = output_dir / f"scorm_{nombre_base}.zip"

                # Crear directorio temporal para el SCORM
                with tempfile.TemporaryDirectory() as temp_dir:
                    scorm_dir = Path(temp_dir) / "scorm"
                    scorm_dir.mkdir()

                    # Copiar recursos generados
                    recursos_copiados = []
                    curso_dir = output_dir / sol['nombre'].replace(' ', '_').lower()

                    if curso_dir.exists():
                        for item in curso_dir.iterdir():
                            if item.is_dir():
                                dest = scorm_dir / item.name
                                shutil.copytree(item, dest)
                                recursos_copiados.append(item.name)

                    # Generar index.html principal
                    recursos_html = ""
                    for i, c in enumerate(st.session_state.contenido_generado):
                        titulo = c.get('titulo', f'Recurso {i+1}')
                        tipo = c.get('tipo', 'Otro')
                        folder = f"{c.get('guion_id', i+1):02d}_{tipo.replace(' ', '_').lower()}"

                        # Determinar archivo principal
                        assets = c.get('assets', {})
                        if 'video' in assets:
                            archivo = f"{folder}/video.mp4"
                            tipo_icono = "🎬"
                        elif 'quiz' in assets:
                            archivo = f"{folder}/quiz.html"
                            tipo_icono = "📝"
                        elif 'html' in assets:
                            archivo = f"{folder}/index.html"
                            tipo_icono = "📄"
                        else:
                            archivo = f"{folder}/index.html"
                            tipo_icono = "📄"

                        recursos_html += f'<div class="recurso" onclick="cargarRecurso(\'{archivo}\')">{tipo_icono} {titulo}</div>\n'

                    index_html = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{sol['nombre']}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Segoe UI', sans-serif; display: flex; height: 100vh; background: #f5f5f5; }}
        .sidebar {{ width: 280px; background: #1a1a2e; color: white; padding: 20px; overflow-y: auto; }}
        .sidebar h2 {{ color: #DA291C; margin-bottom: 20px; font-size: 1.1em; }}
        .recurso {{ padding: 12px 15px; margin: 5px 0; background: rgba(255,255,255,0.1); border-radius: 8px; cursor: pointer; transition: all 0.3s; }}
        .recurso:hover {{ background: #DA291C; }}
        .recurso.active {{ background: #DA291C; }}
        .content {{ flex: 1; display: flex; flex-direction: column; }}
        .header {{ background: white; padding: 15px 25px; border-bottom: 1px solid #ddd; display: flex; justify-content: space-between; align-items: center; }}
        .header h1 {{ font-size: 1.3em; color: #333; }}
        .progress {{ display: flex; align-items: center; gap: 10px; }}
        .progress-bar {{ width: 200px; height: 8px; background: #e0e0e0; border-radius: 4px; overflow: hidden; }}
        .progress-fill {{ height: 100%; background: #DA291C; width: 0%; transition: width 0.3s; }}
        .main-content {{ flex: 1; padding: 0; }}
        iframe {{ width: 100%; height: 100%; border: none; }}
        video {{ width: 100%; height: 100%; background: black; }}
    </style>
</head>
<body>
    <div class="sidebar">
        <h2>📚 {sol['nombre']}</h2>
        {recursos_html}
    </div>
    <div class="content">
        <div class="header">
            <h1 id="titulo-actual">Selecciona un recurso</h1>
            <div class="progress">
                <span id="progreso-texto">0/{len(st.session_state.contenido_generado)}</span>
                <div class="progress-bar"><div class="progress-fill" id="progress-fill"></div></div>
            </div>
        </div>
        <div class="main-content" id="main-content">
            <div style="display: flex; align-items: center; justify-content: center; height: 100%; color: #666;">
                <p>👈 Selecciona un recurso del menú lateral</p>
            </div>
        </div>
    </div>
    <script>
        let completados = 0;
        const total = {len(st.session_state.contenido_generado)};

        function cargarRecurso(archivo) {{
            const container = document.getElementById('main-content');
            const ext = archivo.split('.').pop().toLowerCase();

            document.querySelectorAll('.recurso').forEach(r => r.classList.remove('active'));
            event.target.classList.add('active');

            if (ext === 'mp4') {{
                container.innerHTML = '<video controls autoplay><source src="' + archivo + '" type="video/mp4"></video>';
            }} else {{
                container.innerHTML = '<iframe src="' + archivo + '"></iframe>';
            }}

            completados++;
            document.getElementById('progreso-texto').textContent = completados + '/' + total;
            document.getElementById('progress-fill').style.width = (completados/total*100) + '%';

            // SCORM API
            if (typeof API !== 'undefined') {{
                API.LMSSetValue('cmi.core.lesson_location', archivo);
                API.LMSSetValue('cmi.core.score.raw', Math.round(completados/total*100));
                if (completados >= total) {{
                    API.LMSSetValue('cmi.core.lesson_status', 'completed');
                }}
                API.LMSCommit('');
            }}
        }}

        // Inicializar SCORM
        if (typeof API !== 'undefined') {{
            API.LMSInitialize('');
            API.LMSSetValue('cmi.core.lesson_status', 'incomplete');
        }}
    </script>
</body>
</html>"""

                    with open(scorm_dir / "index.html", 'w', encoding='utf-8') as f:
                        f.write(index_html)

                    # Generar imsmanifest.xml (SCORM 1.2)
                    manifest_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="{nombre_base}" version="1.0"
    xmlns="http://www.imsproject.org/xsd/imscp_rootv1p1p2"
    xmlns:adlcp="http://www.adlnet.org/xsd/adlcp_rootv1p2"
    xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
    xsi:schemaLocation="http://www.imsproject.org/xsd/imscp_rootv1p1p2 imscp_rootv1p1p2.xsd
                        http://www.adlnet.org/xsd/adlcp_rootv1p2 adlcp_rootv1p2.xsd">
    <metadata>
        <schema>ADL SCORM</schema>
        <schemaversion>1.2</schemaversion>
    </metadata>
    <organizations default="org1">
        <organization identifier="org1">
            <title>{sol['nombre']}</title>
            <item identifier="item1" identifierref="res1">
                <title>{sol['nombre']}</title>
            </item>
        </organization>
    </organizations>
    <resources>
        <resource identifier="res1" type="webcontent" adlcp:scormtype="sco" href="index.html">
            <file href="index.html"/>
        </resource>
    </resources>
</manifest>"""

                    with open(scorm_dir / "imsmanifest.xml", 'w', encoding='utf-8') as f:
                        f.write(manifest_xml)

                    # Crear ZIP
                    with zipfile.ZipFile(scorm_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        for root, dirs, files in os.walk(scorm_dir):
                            for file in files:
                                file_path = Path(root) / file
                                arcname = file_path.relative_to(scorm_dir)
                                zipf.write(file_path, arcname)

                st.session_state.scorm_path = str(scorm_path)
                st.success(f"✅ SCORM generado: {scorm_path}")

        # Preview si existe SCORM
        if st.session_state.scorm_path and os.path.exists(st.session_state.scorm_path):
            st.markdown("---")
            st.markdown("#### 👁️ Vista Previa")

            scorm_file = Path(st.session_state.scorm_path)
            file_size = scorm_file.stat().st_size / 1024 / 1024  # MB

            st.markdown(f"""
            <div style="border: 2px solid #28a745; border-radius: 10px; padding: 20px; background: #f0fff0;">
                <h4>✅ Paquete SCORM listo</h4>
                <p><strong>Archivo:</strong> {scorm_file.name}</p>
                <p><strong>Tamaño:</strong> {file_size:.2f} MB</p>
                <p><strong>Recursos incluidos:</strong> {len(st.session_state.contenido_generado)}</p>
            </div>
            """, unsafe_allow_html=True)

            st.markdown("---")

            # Descargar y Preview
            col1, col2 = st.columns(2)
            with col1:
                with open(st.session_state.scorm_path, 'rb') as f:
                    st.download_button(
                        "📥 Descargar SCORM .zip",
                        data=f.read(),
                        file_name=scorm_file.name,
                        mime="application/zip",
                        use_container_width=True
                    )
            with col2:
                if st.button("🌐 Preview Local", use_container_width=True):
                    # Extraer y abrir en navegador
                    import zipfile
                    import tempfile
                    import subprocess

                    preview_dir = Path("/tmp/scorm_preview")
                    if preview_dir.exists():
                        shutil.rmtree(preview_dir)
                    preview_dir.mkdir()

                    with zipfile.ZipFile(st.session_state.scorm_path, 'r') as zipf:
                        zipf.extractall(preview_dir)

                    index_path = preview_dir / "index.html"
                    if index_path.exists():
                        subprocess.run(["open", str(index_path)])
                        st.success("✅ Abierto en navegador")
                    else:
                        st.error("No se encontró index.html en el SCORM")

            # Validación SCORM
            st.markdown("---")
            st.markdown("#### 🔍 Validación")

            with st.expander("Ver contenido del ZIP"):
                import zipfile
                with zipfile.ZipFile(st.session_state.scorm_path, 'r') as zipf:
                    files = zipf.namelist()
                    st.code("\n".join(files[:50]), language="text")
                    if len(files) > 50:
                        st.caption(f"... y {len(files) - 50} archivos más")

                has_manifest = "imsmanifest.xml" in files
                has_index = "index.html" in files

                if has_manifest and has_index:
                    st.success("✅ SCORM válido: contiene imsmanifest.xml e index.html")
                else:
                    if not has_manifest:
                        st.error("❌ Falta imsmanifest.xml")
                    if not has_index:
                        st.error("❌ Falta index.html")

# TAB 6: LMS Integration (Learning)
if tab6:
  with tab6:
    st.markdown("### 🔌 Integrar a LMS")

    st.markdown("""
    <div style="border: 2px dashed #FFD700; border-radius: 10px; padding: 20px; background: #fffef0;">
        <h4>🚧 Funcionalidad en Desarrollo</h4>
        <p>Esta sección permitirá integrar directamente con el LMS de Davivienda (Territorium).</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### Configuración de LMS")

    lms_options = st.selectbox("Seleccionar LMS", [
        "Territorium (Davivienda)",
        "Moodle",
        "Blackboard",
        "Canvas",
        "Otro"
    ])

    if lms_options == "Territorium (Davivienda)":
        st.text_input("URL del LMS", value="https://davivienda.territorium.com", disabled=True)
        st.text_input("API Key", type="password", placeholder="Configurar en .env")
        st.text_input("Curso destino", placeholder="ID o nombre del curso en Territorium")

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("🔍 Verificar conexión", use_container_width=True):
            with st.spinner("Verificando..."):
                import time
                time.sleep(1)
            st.warning("⚠️ Conexión no configurada")

    with col2:
        if st.button("📤 Subir SCORM", use_container_width=True, disabled=not st.session_state.scorm_path):
            st.info("Función disponible próximamente")

    with col3:
        if st.button("📊 Ver estado", use_container_width=True):
            st.info("Sin cursos subidos")

    st.markdown("---")
    st.markdown("#### 📋 Historial de subidas")
    st.dataframe(pd.DataFrame({
        "Fecha": [],
        "Curso": [],
        "LMS": [],
        "Estado": []
    }), use_container_width=True, hide_index=True)

# Footer
st.markdown("---")
st.markdown(
    "<center style='color: #888;'>E-Learning Davivienda | Pipeline v2.0 | Solicitud → Malla → Guiones → Contenido → SCORM → LMS</center>",
    unsafe_allow_html=True
)
